from __future__ import annotations

import json
import random
import os
import uuid
from datetime import datetime

from io import BytesIO
import tempfile
from flask import Blueprint, current_app, jsonify, request, send_file
from sqlalchemy import and_, delete, func, insert, select, update, desc, distinct
from sqlalchemy.exc import SQLAlchemyError

from app.db import get_db, get_session
from docx import Document
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.oxml.ns import qn

try:
    from docx2pdf import convert as docx2pdf_convert
except Exception:
    docx2pdf_convert = None

try:
    import pythoncom
except ImportError:
    pythoncom = None

papers_bp = Blueprint("papers", __name__)


def _table(name: str):
    db = get_db(current_app)
    if name not in db.metadata.tables:
        db.metadata.reflect(bind=db.engine, only=[name])
    return db.metadata.tables[name]


def _parse_strategy(payload: dict) -> dict:
    strategy = payload.get("strategy")
    if isinstance(strategy, str):
        strategy = json.loads(strategy)
    if not isinstance(strategy, dict):
        raise ValueError("strategy 必须是 JSON 对象")

    if not strategy.get("subject_id"):
        raise ValueError("strategy.subject_id 必填")

    sections = strategy.get("sections")
    if not isinstance(sections, list) or not sections:
        raise ValueError("strategy.sections 必须是非空数组")

    normalized = {
        "paper_name": strategy.get("paper_name") or payload.get("paper_name") or "未命名试卷",
        "subject_id": int(strategy["subject_id"]),
        "paper_desc": strategy.get("paper_desc") or "",
        "sections": [],
        "shuffle": bool(strategy.get("shuffle", True)),
    }

    for i, s in enumerate(sections, start=1):
        if not isinstance(s, dict):
            raise ValueError(f"sections[{i}] 必须是对象")
        if not s.get("type_id"):
            raise ValueError(f"sections[{i}].type_id 必填")
        if not s.get("count"):
            raise ValueError(f"sections[{i}].count 必填")

        chapter_ids = s.get("chapter_ids")
        if chapter_ids is None:
            chapter_ids = []
        if not isinstance(chapter_ids, list):
            raise ValueError(f"sections[{i}].chapter_ids 必须是数组")

        normalized["sections"].append(
            {
                "name": s.get("name") or f"第{i}部分",
                "type_id": int(s["type_id"]),
                "difficulty_id": int(s["difficulty_id"]) if s.get("difficulty_id") else None,
                "chapter_ids": [int(x) for x in chapter_ids if x is not None and str(x) != ""],
                "count": int(s["count"]),
                "score_each": float(s["score_each"]) if s.get("score_each") is not None else None,
            }
        )

    return normalized


def _pick_questions(strategy: dict) -> tuple[list[dict], float]:
    qb = _table("question_bank")
    session = get_session(current_app)
    used: set[int] = set()
    picked: list[dict] = []
    total_score = 0.0

    for section in strategy["sections"]:
        where = [qb.c.review_status == 1, qb.c.subject_id == strategy["subject_id"], qb.c.type_id == section["type_id"]]
        if section["difficulty_id"] is not None:
            where.append(qb.c.difficulty_id == section["difficulty_id"])
        if section["chapter_ids"]:
            where.append(qb.c.chapter_id.in_(section["chapter_ids"]))

        candidates_stmt = select(
            qb.c.question_id,
            qb.c.question_content,
            qb.c.question_answer,
            qb.c.question_analysis,
            qb.c.question_score,
            qb.c.chapter_id,
            qb.c.type_id,
            qb.c.difficulty_id,
        ).where(and_(*where))

        candidates = session.execute(candidates_stmt).mappings().all()
        candidates = [dict(c) for c in candidates if int(c["question_id"]) not in used]

        if len(candidates) < section["count"]:
            raise ValueError(f"{section['name']} 可用题目不足：需要{section['count']}，实际{len(candidates)}")

        chosen = random.sample(candidates, section["count"])
        if strategy.get("shuffle", True):
            random.shuffle(chosen)

        for c in chosen:
            qid = int(c["question_id"])
            used.add(qid)
            score = section["score_each"] if section["score_each"] is not None else float(c.get("question_score") or 0)
            total_score += score
            picked.append(
                {
                    "question_id": qid,
                    "question_sort": len(picked) + 1,
                    "question_score": score,
                    "question_content": c.get("question_content"),
                    "question_answer": c.get("question_answer"),
                    "question_analysis": c.get("question_analysis"),
                    "chapter_id": c.get("chapter_id"),
                    "type_id": c.get("type_id"),
                    "difficulty_id": c.get("difficulty_id"),
                    "section_name": section["name"],
                }
            )

    return picked, total_score


@papers_bp.post("/generate")
def generate_preview():
    payload = request.get_json(silent=True) or {}
    try:
        strategy = _parse_strategy(payload)
        questions, total_score = _pick_questions(strategy)
        return jsonify({"paper": {"paper_name": strategy["paper_name"], "subject_id": strategy["subject_id"], "total_score": total_score}, "questions": questions, "strategy": strategy})
    except Exception as err:
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 400


@papers_bp.post("")
def generate_and_save():
    payload = request.get_json(silent=True) or {}
    creator = payload.get("creator") or "creator"
    exam_duration = payload.get("exam_duration")
    is_closed_book = payload.get("is_closed_book")

    try:
        strategy = _parse_strategy(payload)
        questions, total_score = _pick_questions(strategy)
    except Exception as err:
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 400

    paper = _table("exam_paper")
    rel = _table("paper_question_relation")
    session = get_session(current_app)
    now = datetime.now()

    paper_data = {
        "paper_name": strategy["paper_name"],
        "subject_id": strategy["subject_id"],
        "total_score": total_score,
        "exam_duration": exam_duration,
        "is_closed_book": is_closed_book,
        "creator": creator,
        "review_status": 0,
        "paper_desc": strategy.get("paper_desc") or "",
        "create_time": now,
        "update_time": now,
    }

    try:
        res = session.execute(insert(paper).values(**paper_data))
        paper_id = res.inserted_primary_key[0] if res.inserted_primary_key else None
        if not paper_id:
            raise RuntimeError("创建试卷失败")

        for q in questions:
            session.execute(
                insert(rel).values(
                    paper_id=paper_id,
                    question_id=q["question_id"],
                    question_sort=q["question_sort"],
                    question_score=q["question_score"],
                    create_time=now,
                )
            )

        session.commit()
        return jsonify({"paper_id": paper_id, "total_score": total_score, "question_count": len(questions)})
    except SQLAlchemyError as err:
        session.rollback()
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@papers_bp.post("/manual")
def create_manual_paper():
    payload = request.get_json(silent=True) or {}
    creator = payload.get("creator") or "creator"
    exam_duration = payload.get("exam_duration")
    is_closed_book = payload.get("is_closed_book")
    paper_name = payload.get("paper_name") or "未命名试卷"
    paper_desc = payload.get("paper_desc") or ""
    subject_id = payload.get("subject_id")
    items = payload.get("items") or []

    if not subject_id:
        return jsonify({"error": {"message": "subject_id 必填", "type": "BadRequest"}}), 400
    if not isinstance(items, list) or not items:
        return jsonify({"error": {"message": "items 必须是非空数组", "type": "BadRequest"}}), 400

    normalized_items: list[dict] = []
    seen_qids: set[int] = set()
    seen_sorts: set[int] = set()
    for idx, it in enumerate(items, start=1):
        qid = it.get("question_id")
        if not qid:
            return jsonify({"error": {"message": f"items[{idx}].question_id 必填", "type": "BadRequest"}}), 400
        try:
            qid = int(qid)
        except Exception:
            return jsonify({"error": {"message": f"items[{idx}].question_id 格式错误", "type": "BadRequest"}}), 400
        if qid in seen_qids:
            return jsonify({"error": {"message": "题目重复选择", "type": "BadRequest"}}), 400
        seen_qids.add(qid)

        sort = it.get("question_sort")
        if sort is None or sort == "":
            sort = idx
        try:
            sort = int(sort)
        except Exception:
            return jsonify({"error": {"message": f"items[{idx}].question_sort 格式错误", "type": "BadRequest"}}), 400
        if sort < 1:
            return jsonify({"error": {"message": "question_sort 必须 >= 1", "type": "BadRequest"}}), 400
        if sort in seen_sorts:
            return jsonify({"error": {"message": "题号重复，请调整后再保存", "type": "BadRequest"}}), 400
        seen_sorts.add(sort)

        score = it.get("question_score")
        if score is not None and score != "":
            try:
                score = float(score)
            except Exception:
                return jsonify({"error": {"message": f"items[{idx}].question_score 格式错误", "type": "BadRequest"}}), 400
        else:
            score = None

        normalized_items.append({"question_id": qid, "question_sort": sort, "question_score": score})

    qb = _table("question_bank")
    paper = _table("exam_paper")
    rel = _table("paper_question_relation")
    session = get_session(current_app)
    now = datetime.now()

    try:
        q_stmt = (
            select(qb.c.question_id, qb.c.subject_id, qb.c.review_status, qb.c.question_score)
            .where(qb.c.question_id.in_([x["question_id"] for x in normalized_items]))
        )
        rows = session.execute(q_stmt).mappings().all()
        by_id = {int(r["question_id"]): dict(r) for r in rows}
        if len(by_id) != len(normalized_items):
            return jsonify({"error": {"message": "存在无效题目ID", "type": "BadRequest"}}), 400

        total_score = 0.0
        for it in normalized_items:
            r = by_id[it["question_id"]]
            if int(r.get("subject_id") or 0) != int(subject_id):
                return jsonify({"error": {"message": "所选题目科目不一致", "type": "BadRequest"}}), 400
            if int(r.get("review_status") or 0) != 1:
                return jsonify({"error": {"message": "只能选择已通过校验的题目", "type": "BadRequest"}}), 400
            if it["question_score"] is None:
                it["question_score"] = float(r.get("question_score") or 0)
            total_score += float(it["question_score"] or 0)

        paper_data = {
            "paper_name": paper_name,
            "subject_id": int(subject_id),
            "total_score": total_score,
            "exam_duration": exam_duration,
            "is_closed_book": is_closed_book,
            "creator": creator,
            "review_status": 0,
            "paper_desc": paper_desc,
            "create_time": now,
            "update_time": now,
        }

        res = session.execute(insert(paper).values(**paper_data))
        paper_id = res.inserted_primary_key[0] if res.inserted_primary_key else None
        if not paper_id:
            raise RuntimeError("创建试卷失败")

        for it in sorted(normalized_items, key=lambda x: x["question_sort"]):
            session.execute(
                insert(rel).values(
                    paper_id=paper_id,
                    question_id=it["question_id"],
                    question_sort=it["question_sort"],
                    question_score=it["question_score"],
                    create_time=now,
                )
            )

        session.commit()
        return jsonify({"paper_id": paper_id, "total_score": total_score, "question_count": len(normalized_items)})
    except SQLAlchemyError as err:
        session.rollback()
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@papers_bp.get("/reviewers")
def list_reviewers():
    paper = _table("exam_paper")
    try:
        stmt = select(distinct(paper.c.reviewer)).where(and_(paper.c.reviewer != None, paper.c.reviewer != "")).order_by(paper.c.reviewer)
        rows = get_session(current_app).execute(stmt).all()
        return jsonify({"items": [r[0] for r in rows if r[0]]})
    except SQLAlchemyError as err:
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@papers_bp.get("")
def list_papers():
    paper = _table("exam_paper")
    subject_id = request.args.get("subject_id", type=int)
    textbook_id = request.args.get("textbook_id", type=int)
    reviewer = request.args.get("reviewer")
    author = request.args.get("author")
    publisher = request.args.get("publisher")
    review_status = request.args.get("review_status", type=int)

    stmt = select(
        paper.c.paper_id,
        paper.c.paper_name,
        paper.c.subject_id,
        paper.c.total_score,
        paper.c.creator,
        paper.c.review_status,
        paper.c.reviewer,
        paper.c.review_time,
        paper.c.create_time,
        paper.c.update_time,
    )
    if subject_id is not None:
        stmt = stmt.where(paper.c.subject_id == subject_id)
    
    if reviewer:
        stmt = stmt.where(paper.c.reviewer == reviewer)

    if review_status is not None:
        stmt = stmt.where(paper.c.review_status == review_status)
        
    if textbook_id is not None or author or publisher:
        pqr = _table("paper_question_relation")
        qb = _table("question_bank")
        ch = _table("textbook_chapter")
        tb = _table("textbook")
        
        sub = (
            select(1)
            .select_from(pqr)
            .join(qb, qb.c.question_id == pqr.c.question_id)
            .join(ch, ch.c.chapter_id == qb.c.chapter_id)
            .join(tb, tb.c.textbook_id == ch.c.textbook_id)
            .where(pqr.c.paper_id == paper.c.paper_id)
        )
        
        if textbook_id is not None:
            sub = sub.where(ch.c.textbook_id == textbook_id)
        if author:
            sub = sub.where(tb.c.author == author)
        if publisher:
            sub = sub.where(tb.c.publisher == publisher)
            
        stmt = stmt.where(sub.exists())

    stmt = stmt.order_by(paper.c.paper_id.desc()).limit(200)
    try:
        rows = get_session(current_app).execute(stmt).mappings().all()
        return jsonify({"items": [dict(r) for r in rows]})
    except SQLAlchemyError as err:
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@papers_bp.get("/<int:paper_id>")
def get_paper(paper_id: int):
    paper = _table("exam_paper")
    rel = _table("paper_question_relation")
    qb = _table("question_bank")
    session = get_session(current_app)
    try:
        p = session.execute(select(paper).where(paper.c.paper_id == paper_id)).mappings().first()
        if p is None:
            return jsonify({"error": {"message": "试卷不存在", "type": "NotFound"}}), 404

        q_stmt = (
            select(
                rel.c.question_sort,
                rel.c.question_score,
                qb.c.question_id,
                qb.c.question_content,
                qb.c.question_answer,
                qb.c.question_analysis,
                qb.c.type_id,
                qb.c.difficulty_id,
                qb.c.chapter_id,
            )
            .join(qb, qb.c.question_id == rel.c.question_id)
            .where(rel.c.paper_id == paper_id)
            .order_by(rel.c.question_sort.asc())
        )
        questions = session.execute(q_stmt).mappings().all()
        return jsonify({"paper": dict(p), "questions": [dict(r) for r in questions]})
    except SQLAlchemyError as err:
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@papers_bp.delete("/<int:paper_id>")
def delete_paper(paper_id: int):
    paper = _table("exam_paper")
    rel = _table("paper_question_relation")
    session = get_session(current_app)
    try:
        session.execute(delete(rel).where(rel.c.paper_id == paper_id))
        session.execute(delete(paper).where(paper.c.paper_id == paper_id))
        session.commit()
        return jsonify({"ok": True})
    except SQLAlchemyError as err:
        session.rollback()
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@papers_bp.post("/batch-delete")
def batch_delete_papers():
    payload = request.get_json(silent=True) or {}
    ids = payload.get("ids")
    if not isinstance(ids, list) or not ids:
        return jsonify({"error": {"message": "ids 必须是非空数组", "type": "BadRequest"}}), 400

    paper = _table("exam_paper")
    rel = _table("paper_question_relation")
    session = get_session(current_app)
    try:
        session.execute(delete(rel).where(rel.c.paper_id.in_(ids)))
        session.execute(delete(paper).where(paper.c.paper_id.in_(ids)))
        session.commit()
        return jsonify({"ok": True})
    except SQLAlchemyError as err:
        session.rollback()
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@papers_bp.put("/<int:paper_id>")
def update_paper(paper_id: int):
    payload = request.get_json(silent=True) or {}
    paper = _table("exam_paper")
    fields = {}
    for k in ["paper_name", "exam_duration", "is_closed_book", "paper_desc", "review_status", "reviewer", "review_time", "total_score"]:
        if k in payload:
            fields[k] = payload.get(k)
    fields["update_time"] = datetime.now()
    try:
        session = get_session(current_app)
        session.execute(update(paper).where(paper.c.paper_id == paper_id).values(**fields))
        session.commit()
        return jsonify({"ok": True})
    except SQLAlchemyError as err:
        session.rollback()
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@papers_bp.put("/<int:paper_id>/questions")
def update_paper_questions(paper_id: int):
    payload = request.get_json(silent=True) or {}
    items = payload.get("items") or []
    if not isinstance(items, list) or not items:
        return jsonify({"error": {"message": "items 必须是非空数组", "type": "BadRequest"}}), 400

    rel = _table("paper_question_relation")
    session = get_session(current_app)
    try:
        for it in items:
            if not it.get("question_id"):
                continue
            data = {}
            if "question_sort" in it:
                data["question_sort"] = int(it["question_sort"])
            if "question_score" in it:
                data["question_score"] = it["question_score"]
            if data:
                session.execute(
                    update(rel)
                    .where(and_(rel.c.paper_id == paper_id, rel.c.question_id == int(it["question_id"])))
                    .values(**data)
                )
        session.commit()
        return jsonify({"ok": True})
    except SQLAlchemyError as err:
        session.rollback()
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


def _export_base_dir(paper_id: int) -> str:
    base = os.path.join(current_app.config["EXPORT_DIR"], "papers", str(paper_id))
    os.makedirs(base, exist_ok=True)
    return base


@papers_bp.get("/<int:paper_id>/exports")
def list_exports(paper_id: int):
    history = _table("paper_export_history")
    session = get_session(current_app)
    try:
        stmt = select(history).where(history.c.paper_id == paper_id).order_by(desc(history.c.created_at))
        rows = session.execute(stmt).mappings().all()
        return jsonify({"items": [dict(r) for r in rows]})
    except SQLAlchemyError as err:
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@papers_bp.get("/<int:paper_id>/exports/<string:version_id>/download")
def download_export(paper_id: int, version_id: str):
    history = _table("paper_export_history")
    session = get_session(current_app)
    
    try:
        stmt = select(history).where(and_(history.c.paper_id == paper_id, history.c.version_id == version_id))
        row = session.execute(stmt).mappings().first()
        
        if not row:
            return jsonify({"error": {"message": "版本不存在", "type": "NotFound"}}), 404
            
        path = os.path.join(_export_base_dir(paper_id), row["filename"])
        if not os.path.exists(path):
            return jsonify({"error": {"message": "文件不存在", "type": "NotFound"}}), 404
            
        return send_file(path, as_attachment=True, download_name=row["download_name"] or row["filename"])
    except SQLAlchemyError as err:
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


def _set_columns(section, cols):
    sectPr = section._sectPr
    cols_elm = sectPr.xpath('./w:cols')
    if cols_elm:
        cols_elm = cols_elm[0]
    else:
        cols_elm = OxmlElement('w:cols')
        sectPr.append(cols_elm)

    cols_elm.set(qn('w:num'), str(cols))
    cols_elm.set(qn('w:space'), '425') # ~1.5cm gap
    cols_elm.set(qn('w:sep'), '1') # Separator line

def _render_word(paper_row: dict, questions: list[dict], header: str | None, footer: str | None, include_answer: bool, paper_size: str = 'A4') -> Document:
    doc = Document()
    section = doc.sections[0]
    
    if paper_size == 'A3':
        section.page_width = Mm(420)
        section.page_height = Mm(297)
        section.orientation = WD_ORIENT.LANDSCAPE
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        
        # Add title spanning columns
        # In Word, to span columns, we usually put title in a separate section (1 col)
        # then content in another section (2 cols).
        # doc.sections[0] is already there. Let's use it for Title.
        
        p = doc.add_paragraph(paper_row.get("paper_name") or "")
        p.alignment = 1 # Center
        run = p.runs[0]
        run.font.size = Pt(18)
        run.bold = True
        
        if paper_row.get("paper_desc"):
            doc.add_paragraph(str(paper_row["paper_desc"])).alignment = 1
            
        # Meta info
        meta = f"时长：{paper_row.get('exam_duration')}分钟    {'闭卷' if paper_row.get('is_closed_book') else '开卷'}    总分：{paper_row.get('total_score')}"
        doc.add_paragraph(meta).alignment = 1
        
        doc.add_paragraph() # Spacer
        
        # Start new section for columns
        new_section = doc.add_section(WD_SECTION.CONTINUOUS)
        _set_columns(new_section, 2)
        
    else:
        # A4 Default
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.orientation = WD_ORIENT.PORTRAIT
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        
        doc.add_paragraph(paper_row.get("paper_name") or "").alignment = 1
        if paper_row.get("paper_desc"):
            doc.add_paragraph(str(paper_row["paper_desc"]))
        
        meta = f"时长：{paper_row.get('exam_duration')}分钟    {'闭卷' if paper_row.get('is_closed_book') else '开卷'}    总分：{paper_row.get('total_score')}"
        doc.add_paragraph(meta).alignment = 1
        doc.add_paragraph()

    if header:
        doc.add_paragraph(str(header))

    for q in questions:
        p = doc.add_paragraph()
        p.add_run(f"{q['question_sort']}. ").bold = True
        if q.get('question_score') is not None:
            p.add_run(f"（{q.get('question_score')}分） ")
        p.add_run(q.get('question_content') or '')
        
        if include_answer:
            if q.get("question_answer"):
                doc.add_paragraph(f"答案：{q.get('question_answer')}")
            if q.get("question_analysis"):
                doc.add_paragraph(f"解析：{q.get('question_analysis')}")
        
        doc.add_paragraph() # Spacer

    if footer:
        doc.add_paragraph(str(footer))
    return doc


@papers_bp.post("/<int:paper_id>/export/word")
def export_word(paper_id: int):
    payload = request.get_json(silent=True) or {}
    header = payload.get("header")
    footer = payload.get("footer")
    include_answer = bool(payload.get("include_answer", False))
    paper_size = payload.get("paper_size", "A4") # A4 or A3
    payload_questions = payload.get("questions")

    session = get_session(current_app)
    paper = _table("exam_paper")
    rel = _table("paper_question_relation")
    qb = _table("question_bank")

    try:
        p = session.execute(select(paper).where(paper.c.paper_id == paper_id)).mappings().first()
        if p is None:
            return jsonify({"error": {"message": "试卷不存在", "type": "NotFound"}}), 404

        q_stmt = (
            select(
                rel.c.question_sort,
                rel.c.question_score,
                qb.c.question_id,
                qb.c.question_content,
                qb.c.question_answer,
                qb.c.question_analysis,
            )
            .join(qb, qb.c.question_id == rel.c.question_id)
            .where(rel.c.paper_id == paper_id)
            .order_by(rel.c.question_sort.asc())
        )
        questions = [dict(r) for r in session.execute(q_stmt).mappings().all()]
        if isinstance(payload_questions, list) and payload_questions:
            normalized = []
            for idx, q in enumerate(payload_questions):
                if not isinstance(q, dict):
                    continue
                normalized.append({
                    "question_sort": q.get("question_sort") or idx + 1,
                    "question_score": q.get("question_score"),
                    "question_id": q.get("question_id"),
                    "question_content": q.get("question_content") or "",
                    "question_answer": q.get("question_answer") or "",
                    "question_analysis": q.get("question_analysis") or ""
                })
            if normalized:
                questions = normalized

        doc = _render_word(dict(p), questions, header, footer, include_answer, paper_size)

        f = BytesIO()
        doc.save(f)
        f.seek(0)
        
        filename = f"{p['paper_name']}.docx"
        try:
            from urllib.parse import quote
            filename = quote(filename)
        except:
            pass

        return send_file(f, as_attachment=True, download_name=f"{p['paper_name']}.docx")
    except SQLAlchemyError as err:
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@papers_bp.post("/<int:paper_id>/export/pdf")
def export_pdf(paper_id: int):
    if docx2pdf_convert is None:
        return jsonify({"error": {"message": "PDF 导出不可用：docx2pdf 未安装或不可用", "type": "NotSupported"}}), 400

    payload = request.get_json(silent=True) or {}
    header = payload.get("header")
    footer = payload.get("footer")
    include_answer = bool(payload.get("include_answer", False))
    paper_size = payload.get("paper_size", "A4")
    payload_questions = payload.get("questions")

    session = get_session(current_app)
    paper = _table("exam_paper")
    rel = _table("paper_question_relation")
    qb = _table("question_bank")

    try:
        p = session.execute(select(paper).where(paper.c.paper_id == paper_id)).mappings().first()
        if p is None:
            return jsonify({"error": {"message": "试卷不存在", "type": "NotFound"}}), 404

        q_stmt = (
            select(
                rel.c.question_sort,
                rel.c.question_score,
                qb.c.question_id,
                qb.c.question_content,
                qb.c.question_answer,
                qb.c.question_analysis,
            )
            .join(qb, qb.c.question_id == rel.c.question_id)
            .where(rel.c.paper_id == paper_id)
            .order_by(rel.c.question_sort.asc())
        )
        questions = [dict(r) for r in session.execute(q_stmt).mappings().all()]
        if isinstance(payload_questions, list) and payload_questions:
            normalized = []
            for idx, q in enumerate(payload_questions):
                if not isinstance(q, dict):
                    continue
                normalized.append({
                    "question_sort": q.get("question_sort") or idx + 1,
                    "question_score": q.get("question_score"),
                    "question_id": q.get("question_id"),
                    "question_content": q.get("question_content") or "",
                    "question_answer": q.get("question_answer") or "",
                    "question_analysis": q.get("question_analysis") or ""
                })
            if normalized:
                questions = normalized

        doc = _render_word(dict(p), questions, header, footer, include_answer, paper_size)

        # Use temp files for conversion
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
            doc.save(tmp_docx.name)
            tmp_docx_path = os.path.abspath(tmp_docx.name)
            
        tmp_pdf_path = tmp_docx_path.replace(".docx", ".pdf")
        
        try:
            if pythoncom:
                pythoncom.CoInitialize()
            
            docx2pdf_convert(tmp_docx_path, tmp_pdf_path)
            
            with open(tmp_pdf_path, "rb") as f:
                pdf_data = f.read()
                
            return send_file(
                BytesIO(pdf_data), 
                as_attachment=True, 
                download_name=f"{p['paper_name']}.pdf",
                mimetype='application/pdf'
            )
        finally:
            if pythoncom:
                pythoncom.CoUninitialize()
            
            if os.path.exists(tmp_docx_path):
                try:
                    os.remove(tmp_docx_path)
                except:
                    pass
            if os.path.exists(tmp_pdf_path):
                try:
                    os.remove(tmp_pdf_path)
                except:
                    pass
    except Exception as e:
        print(f"PDF Export Error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": {"message": f"PDF转换失败: {str(e)}", "type": "ConversionError"}}), 500
