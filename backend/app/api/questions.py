from __future__ import annotations

import io
import re
from decimal import Decimal

from flask import Blueprint, current_app, jsonify, request
from sqlalchemy import and_, delete, func, insert, or_, select, update
from sqlalchemy.exc import SQLAlchemyError

from app.db import get_db, get_session
from docx import Document
from openpyxl import load_workbook

questions_bp = Blueprint("questions", __name__)


def _table(name: str):
    db = get_db(current_app)
    if name not in db.metadata.tables:
        db.metadata.reflect(bind=db.engine, only=[name])
    return db.metadata.tables[name]


def _to_jsonable(v):
    if isinstance(v, Decimal):
        return float(v)
    return v


def _int_or_none(v):
    if v is None or v == "":
        return None
    try:
        return int(v)
    except Exception:
        return None


def _decimal_or_none(v):
    if v is None or v == "":
        return None
    try:
        return Decimal(str(v))
    except Exception:
        return None


def _parse_csv_ints(v: str | None) -> list[int]:
    if v is None:
        return []
    if isinstance(v, str):
        parts = [x.strip() for x in v.split(",") if x.strip()]
        out: list[int] = []
        for p in parts:
            try:
                out.append(int(p))
            except Exception:
                continue
        return out
    return []


@questions_bp.get("")
def search_questions():
    q = request.args.get("q", type=str)
    subject_id = request.args.get("subject_id", type=int)
    # chapter_id = request.args.get("chapter_id", type=int) # 旧逻辑
    chapter_ids_str = request.args.get("chapter_id", type=str)
    ids_str = request.args.get("ids", type=str)

    type_ids_str = request.args.get("type_id", type=str)
    difficulty_ids_str = request.args.get("difficulty_id", type=str)
    textbook_id = request.args.get("textbook_id", type=int)
    author = request.args.get("author", type=str)
    publisher = request.args.get("publisher", type=str)
    review_status = request.args.get("review_status", type=int)

    page = max(1, request.args.get("page", default=1, type=int))
    page_size = min(100, max(1, request.args.get("page_size", default=20, type=int)))

    t = _table("question_bank")
    ch = _table("textbook_chapter")
    tb = _table("textbook")
    sd = _table("subject_dict")
    qtd = _table("question_type_dict")
    qdd = _table("question_difficulty_dict")

    from_ = (
        t.outerjoin(ch, ch.c.chapter_id == t.c.chapter_id)
        .outerjoin(tb, tb.c.textbook_id == ch.c.textbook_id)
        .outerjoin(sd, sd.c.subject_id == t.c.subject_id)
        .outerjoin(qtd, qtd.c.type_id == t.c.type_id)
        .outerjoin(qdd, qdd.c.difficulty_id == t.c.difficulty_id)
    )

    where = []
    if subject_id is not None:
        where.append(t.c.subject_id == subject_id)
    
    # 新增：处理多选章节
    if chapter_ids_str:
        # 尝试逗号分割
        try:
            c_ids = [int(x) for x in chapter_ids_str.split(",") if x.strip()]
            if c_ids:
                if len(c_ids) == 1:
                    where.append(t.c.chapter_id == c_ids[0])
                else:
                    where.append(t.c.chapter_id.in_(c_ids))
        except ValueError:
            pass

    # 新增：处理 IDs 筛选
    if ids_str:
        try:
            ids = [int(x) for x in ids_str.split(",") if x.strip()]
            if ids:
                if len(ids) == 1:
                    where.append(t.c.question_id == ids[0])
                else:
                    where.append(t.c.question_id.in_(ids))
        except ValueError:
            pass

    type_ids = _parse_csv_ints(type_ids_str)
    if type_ids:
        if len(type_ids) == 1:
            where.append(t.c.type_id == type_ids[0])
        else:
            where.append(t.c.type_id.in_(type_ids))

    difficulty_ids = _parse_csv_ints(difficulty_ids_str)
    if difficulty_ids:
        if len(difficulty_ids) == 1:
            where.append(t.c.difficulty_id == difficulty_ids[0])
        else:
            where.append(t.c.difficulty_id.in_(difficulty_ids))

    if textbook_id is not None:
        where.append(tb.c.textbook_id == textbook_id)
    if author:
        where.append(tb.c.author.like(f"%{author}%"))
    if publisher:
        where.append(tb.c.publisher.like(f"%{publisher}%"))
    if review_status is not None:
        where.append(t.c.review_status == review_status)
    if q:
        like = f"%{q}%"
        where.append(or_(t.c.question_content.like(like), t.c.question_analysis.like(like)))

    stmt = (
        select(
            t.c.question_id,
            t.c.subject_id,
            t.c.chapter_id,
            t.c.type_id,
            t.c.difficulty_id,
            sd.c.subject_name.label("subject_name"),
            ch.c.chapter_name.label("chapter_name"),
            qtd.c.type_name.label("type_name"),
            qdd.c.difficulty_name.label("difficulty_name"),
            tb.c.textbook_id.label("textbook_id"),
            tb.c.textbook_name.label("textbook_name"),
            tb.c.author.label("author"),
            tb.c.publisher.label("publisher"),
            t.c.question_content,
            t.c.question_answer,
            t.c.question_analysis,
            t.c.question_score,
            t.c.is_ai_generated,
            t.c.source_question_ids,
            t.c.reviewer,
            t.c.review_time,
            t.c.review_status,
            t.c.create_user,
            t.c.create_time,
            t.c.update_time,
        )
        .select_from(from_)
        .where(and_(*where) if where else True)
        .order_by(t.c.question_id.desc())
        .offset((page - 1) * page_size)
        .limit(page_size)
    )

    count_stmt = select(func.count()).select_from(from_).where(and_(*where) if where else True)

    try:
        session = get_session(current_app)
        total = session.execute(count_stmt).scalar_one()
        rows = session.execute(stmt).mappings().all()
        items = [{k: _to_jsonable(v) for k, v in dict(r).items()} for r in rows]
        return jsonify({"items": items, "total": total, "page": page, "page_size": page_size})
    except SQLAlchemyError as err:
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@questions_bp.post("")
def create_question():
    payload = request.get_json(silent=True) or {}
    t = _table("question_bank")

    required = ["subject_id", "chapter_id", "type_id", "difficulty_id", "question_content"]
    missing = [k for k in required if payload.get(k) in [None, ""]]
    if missing:
        return jsonify({"error": {"message": f"缺少字段: {', '.join(missing)}", "type": "BadRequest"}}), 400

    score = payload.get("question_score")
    if score is not None:
        try:
            score = Decimal(str(score))
        except Exception:
            return jsonify({"error": {"message": "question_score 格式错误", "type": "BadRequest"}}), 400

    data = {
        "subject_id": payload.get("subject_id"),
        "chapter_id": payload.get("chapter_id"),
        "type_id": payload.get("type_id"),
        "difficulty_id": payload.get("difficulty_id"),
        "question_content": payload.get("question_content"),
        "question_answer": payload.get("question_answer"),
        "question_analysis": payload.get("question_analysis"),
        "question_score": score,
        "is_ai_generated": 0,
        "source_question_ids": None,
        "review_status": 1,
        "reviewer": payload.get("create_user") or "manual",
        "create_user": payload.get("create_user") or "manual",
    }

    try:
        session = get_session(current_app)
        res = session.execute(insert(t).values(**data))
        session.commit()
        question_id = res.inserted_primary_key[0] if res.inserted_primary_key else None
        return jsonify({"question_id": question_id})
    except SQLAlchemyError as err:
        get_session(current_app).rollback()
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@questions_bp.get("/<int:question_id>")
def get_question(question_id: int):
    t = _table("question_bank")
    stmt = select(t).where(t.c.question_id == question_id)
    try:
        row = get_session(current_app).execute(stmt).mappings().first()
        if row is None:
            return jsonify({"error": {"message": "题目不存在", "type": "NotFound"}}), 404
        item = {k: _to_jsonable(v) for k, v in dict(row).items()}
        return jsonify({"item": item})
    except SQLAlchemyError as err:
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@questions_bp.put("/<int:question_id>")
def update_question(question_id: int):
    payload = request.get_json(silent=True) or {}
    t = _table("question_bank")

    data = {}
    for k in [
        "subject_id",
        "chapter_id",
        "type_id",
        "difficulty_id",
        "question_content",
        "question_answer",
        "question_analysis",
        "review_status",
        "reviewer",
        "review_time",
    ]:
        if k in payload:
            data[k] = payload.get(k)

    if "question_score" in payload:
        score = payload.get("question_score")
        if score is None:
            data["question_score"] = None
        else:
            try:
                data["question_score"] = Decimal(str(score))
            except Exception:
                return jsonify({"error": {"message": "question_score 格式错误", "type": "BadRequest"}}), 400

    if not data:
        return jsonify({"ok": True})

    try:
        session = get_session(current_app)
        session.execute(update(t).where(t.c.question_id == question_id).values(**data))
        session.commit()
        return jsonify({"ok": True})
    except SQLAlchemyError as err:
        get_session(current_app).rollback()
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@questions_bp.delete("/<int:question_id>")
def delete_question(question_id: int):
    t = _table("question_bank")
    try:
        session = get_session(current_app)
        session.execute(delete(t).where(t.c.question_id == question_id))
        session.commit()
        return jsonify({"ok": True})
    except SQLAlchemyError as err:
        get_session(current_app).rollback()
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@questions_bp.post("/import/excel")
def import_questions_excel():
    file = request.files.get("file")
    if file is None:
        return jsonify({"error": {"message": "缺少上传文件 file", "type": "BadRequest"}}), 400

    default_subject_id = _int_or_none(request.form.get("subject_id"))
    default_chapter_id = _int_or_none(request.form.get("chapter_id"))
    default_type_id = _int_or_none(request.form.get("type_id"))
    default_difficulty_id = _int_or_none(request.form.get("difficulty_id"))
    create_user = request.form.get("create_user") or "import"

    wb = load_workbook(filename=io.BytesIO(file.read()), data_only=True)
    ws = wb.active

    header = [str(c.value).strip() if c.value is not None else "" for c in next(ws.iter_rows(min_row=1, max_row=1))[0:]]
    idx = {h: i for i, h in enumerate(header) if h}

    def pick(row, *names):
        for n in names:
            if n in idx:
                v = row[idx[n]].value
                if v is not None and str(v).strip() != "":
                    return v
        return None

    t = _table("question_bank")
    session = get_session(current_app)
    inserted = 0
    skipped = 0

    try:
        for row in ws.iter_rows(min_row=2):
            content = pick(row, "question_content", "题干", "题目内容")
            if content is None or str(content).strip() == "":
                skipped += 1
                continue

            subject_id = _int_or_none(pick(row, "subject_id", "科目ID")) or default_subject_id
            chapter_id = _int_or_none(pick(row, "chapter_id", "章节ID")) or default_chapter_id
            type_id = _int_or_none(pick(row, "type_id", "题型ID")) or default_type_id
            difficulty_id = _int_or_none(pick(row, "difficulty_id", "难度ID")) or default_difficulty_id
            if not subject_id or not chapter_id or not type_id or not difficulty_id:
                skipped += 1
                continue

            data = {
                "subject_id": subject_id,
                "chapter_id": chapter_id,
                "type_id": type_id,
                "difficulty_id": difficulty_id,
                "question_content": str(content).strip(),
                "question_answer": pick(row, "question_answer", "答案"),
                "question_analysis": pick(row, "question_analysis", "解析"),
                "question_score": _decimal_or_none(pick(row, "question_score", "分值")),
                "is_ai_generated": 0,
                "source_question_ids": None,
                "review_status": 1,
                "reviewer": create_user,
                "create_user": create_user,
            }
            session.execute(insert(t).values(**data))
            inserted += 1

        session.commit()
        return jsonify({"ok": True, "inserted": inserted, "skipped": skipped})
    except SQLAlchemyError as err:
        session.rollback()
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@questions_bp.post("/import/word")
def import_questions_word():
    file = request.files.get("file")
    if file is None:
        return jsonify({"error": {"message": "缺少上传文件 file", "type": "BadRequest"}}), 400

    default_subject_id = _int_or_none(request.form.get("subject_id"))
    default_chapter_id = _int_or_none(request.form.get("chapter_id"))
    default_type_id = _int_or_none(request.form.get("type_id"))
    default_difficulty_id = _int_or_none(request.form.get("difficulty_id"))
    create_user = request.form.get("create_user") or "import"

    doc = Document(io.BytesIO(file.read()))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

    blocks = []
    cur = {}

    def flush():
        nonlocal cur
        if cur.get("question_content"):
            blocks.append(cur)
        cur = {}

    key_map = {
        "题干": "question_content",
        "题目": "question_content",
        "答案": "question_answer",
        "解析": "question_analysis",
        "分值": "question_score",
        "科目ID": "subject_id",
        "章节ID": "chapter_id",
        "题型ID": "type_id",
        "难度ID": "difficulty_id",
    }

    for line in paragraphs:
        m = re.match(r"^([^:：]{1,8})[:：](.*)$", line)
        if m:
            k = m.group(1).strip()
            v = m.group(2).strip()
            if k in ["题干", "题目"]:
                flush()
            field = key_map.get(k)
            if field:
                if field in ["subject_id", "chapter_id", "type_id", "difficulty_id"]:
                    cur[field] = _int_or_none(v)
                elif field == "question_score":
                    cur[field] = _decimal_or_none(v)
                else:
                    cur[field] = v
                continue

        if "question_content" not in cur:
            cur["question_content"] = line
        else:
            cur["question_content"] = f"{cur['question_content']}\n{line}"

    flush()

    t = _table("question_bank")
    session = get_session(current_app)
    inserted = 0
    skipped = 0

    try:
        for b in blocks:
            subject_id = b.get("subject_id") or default_subject_id
            chapter_id = b.get("chapter_id") or default_chapter_id
            type_id = b.get("type_id") or default_type_id
            difficulty_id = b.get("difficulty_id") or default_difficulty_id
            if not subject_id or not chapter_id or not type_id or not difficulty_id:
                skipped += 1
                continue

            data = {
                "subject_id": subject_id,
                "chapter_id": chapter_id,
                "type_id": type_id,
                "difficulty_id": difficulty_id,
                "question_content": b.get("question_content"),
                "question_answer": b.get("question_answer"),
                "question_analysis": b.get("question_analysis"),
                "question_score": b.get("question_score"),
                "is_ai_generated": 0,
                "source_question_ids": None,
                "review_status": 1,
                "reviewer": create_user,
                "create_user": create_user,
            }
            session.execute(insert(t).values(**data))
            inserted += 1

        session.commit()
        return jsonify({"ok": True, "inserted": inserted, "skipped": skipped})
    except SQLAlchemyError as err:
        session.rollback()
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500
