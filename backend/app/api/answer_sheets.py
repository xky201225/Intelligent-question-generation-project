from __future__ import annotations

import json
import os
import uuid
from datetime import datetime

from io import BytesIO
import tempfile
from flask import Blueprint, current_app, jsonify, request, send_file
from sqlalchemy import and_, delete, func, insert, select, update
from sqlalchemy.exc import IntegrityError, SQLAlchemyError
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.db import get_db, get_session

try:
    from docx2pdf import convert as docx2pdf_convert
except Exception:
    docx2pdf_convert = None

import traceback

answer_styles_bp = Blueprint("answer_styles", __name__)
answer_sheets_bp = Blueprint("answer_sheets", __name__)


def _table(name: str):
    db = get_db(current_app)
    if name not in db.metadata.tables:
        db.metadata.reflect(bind=db.engine, only=[name])
    return db.metadata.tables[name]


def _export_base_dir(sub_dir: str | int) -> str:
    base = current_app.config.get("EXPORT_DIR") or "exports"
    path = os.path.join(base, str(sub_dir))
    if not os.path.exists(path):
        os.makedirs(path)
    return path


def _set_columns(section, cols):
    sectPr = section._sectPr
    cols_elm = sectPr.xpath('./w:cols')
    if cols_elm:
        cols_elm = cols_elm[0]
    else:
        cols_elm = OxmlElement('w:cols')
        sectPr.append(cols_elm)
    
    cols_elm.set(qn('w:num'), str(cols))
    cols_elm.set(qn('w:space'), '425') # ~1.5cm gap
    cols_elm.set(qn('w:sep'), '1') # Separator line

def _render_sheet_word(sheet: dict, items: list[dict], styles: dict[int, dict], paper_size: str = 'A3', ticket_no_digits: int = 10) -> Document:
    doc = Document()
    
    # 1. Page Setup
    section = doc.sections[0]
    if paper_size == 'A4':
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.orientation = WD_ORIENT.PORTRAIT
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
    else:
        # Default A3 Landscape
        section.page_width = Mm(420)
        section.page_height = Mm(297)
        section.orientation = WD_ORIENT.LANDSCAPE
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
    
    # 2. Header Table (Student Info & Ticket Number)
    # Create a table for the header layout
    # Left: Info, Right: Ticket Number Grid
    header_table = doc.add_table(rows=1, cols=2)
    header_table.style = 'Table Grid'
    header_table.autofit = False
    
    # Adjust column widths
    # Total width approx 39cm. Split 60% / 40%
    header_table.columns[0].width = Cm(24)
    header_table.columns[1].width = Cm(15)
    
    # Left Cell: Title + Info
    left_cell = header_table.cell(0, 0)
    left_p = left_cell.paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = left_p.add_run(sheet.get("sheet_name") or "答题卡")
    run.font.size = Pt(18)
    run.bold = True
    
    left_cell.add_paragraph() # Spacer
    
    info_table = left_cell.add_table(rows=3, cols=2)
    info_table.style = None # No border
    # Fill info placeholders
    info_table.cell(0, 0).text = "姓名：__________________"
    info_table.cell(0, 1).text = "班级：__________________"
    info_table.cell(1, 0).text = "考号：__________________"
    info_table.cell(1, 1).text = "座位：__________________"
    info_table.cell(2, 0).text = "注意事项："
    
    note_p = left_cell.add_paragraph()
    note_p.add_run("1. 答题前请将姓名、班级、考号填写清楚。\n").font.size = Pt(9)
    note_p.add_run("2. 客观题必须使用2B铅笔填涂；主观题必须使用黑色签字笔书写。\n").font.size = Pt(9)
    note_p.add_run("3. 必须在各题目的答题区域内作答，超出黑色矩形边框限定区域的答案无效。").font.size = Pt(9)
    
    # Right Cell: Ticket Number Grid (Mockup)
    right_cell = header_table.cell(0, 1)
    right_p = right_cell.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    right_p.add_run("准考证号填涂区").bold = True
    
    # Mock grid for ticket number (10 cols x 10 rows digits)
    # ticket_no_digits comes from request
    ticket_table = right_cell.add_table(rows=10, cols=ticket_no_digits)
    ticket_table.style = 'Table Grid'
    for r in range(10):
        for c in range(ticket_no_digits):
            cell = ticket_table.cell(r, c)
            cell.text = f"[{r}]"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.size = Pt(8)

    doc.add_paragraph() # Spacer
    
    # 3. Content Columns
    # Insert a section break to start columns
    # Only for A3
    if paper_size == 'A3':
        new_section = doc.add_section(WD_SECTION.CONTINUOUS)
        try:
            _set_columns(new_section, 2) # 2 Columns layout for A3
        except Exception:
             # Fallback if XML manip fails, just continue single col? 
             # Or log it.
             pass
    
    # Group items by type to separate Objective (Choice) vs Subjective
    # 1, 7, 2 are usually Objective (Choice/Judge)
    # Others are Subjective
    
    sorted_items = sorted(items, key=lambda x: (x["area_sort"] if x["area_sort"] is not None else 99999))
    
    obj_items = [it for it in sorted_items if styles.get(it["style_id"], {}).get("type_id") in [1, 2, 7]]
    sub_items = [it for it in sorted_items if styles.get(it["style_id"], {}).get("type_id") not in [1, 2, 7]]
    
    # --- Part I: Objective Questions ---
    if obj_items:
        p = doc.add_paragraph()
        p.add_run("第Ⅰ卷（客观题）").bold = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        
        # Render Choice Grid
        # Group by 5 items for better layout
        chunk_size = 5
        for i in range(0, len(obj_items), chunk_size):
            chunk = obj_items[i:i+chunk_size]
            
            # Create a small table for this row of 5 questions
            # Each question takes a cell
            # But wait, we are in a column, width is small (approx 39/3 = 13cm)
            # Maybe just list them vertically or 2 per row in the column?
            # Standard answer sheet usually packs them tight.
            # Let's try 1 question per line but compact
            
            for item in chunk:
                style = styles.get(item["style_id"], {})
                config_str = style.get("style_config") or "{}"
                try: config = json.loads(config_str)
                except: config = {}
                
                try:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(2)
                    p.add_run(f"{item['area_sort']}.").bold = True
                    
                    options = config.get("options") or ["A", "B", "C", "D"]
                    run_text = " "
                    for opt in options:
                        run_text += f"[{opt}] "
                    p.add_run(run_text).font.color.rgb = RGBColor(255, 0, 0) # Red for bubble simulation
                except Exception as e:
                    print(f"Error rendering item {item.get('question_id')}: {e}")
                    p.add_run(" [渲染错误] ").font.color.rgb = RGBColor(255, 0, 0)
    
    # --- Part II: Subjective Questions ---
    if sub_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.add_run("第Ⅱ卷（主观题）").bold = True
        p.add_run("\n注意：请在各题目的答题区域内作答，超出黑色矩形边框限定区域的答案无效。").font.size = Pt(8)
        
        for item in sub_items:
            try:
                style = styles.get(item["style_id"], {})
                config_str = style.get("style_config") or "{}"
                try: config = json.loads(config_str)
                except: config = {}
                
                # Frame for each question
                # Use a single-cell table to create the "black rectangle frame"
                frame_table = doc.add_table(rows=1, cols=1)
                frame_table.style = 'Table Grid'
                frame_table.autofit = False
                
                # Dynamic width based on column layout
                # A3 (2 cols): (42 - 3 margins)/2 = 19.5. Safe ~18.0
                # A4 (1 col): 21 - 4 margins = 17. Safe ~17
                f_width = Cm(18.0) if paper_size == 'A3' else Cm(17)
                frame_table.columns[0].width = f_width
                
                cell = frame_table.cell(0, 0)
                cp = cell.paragraphs[0]
                cp.add_run(f"{item['area_sort']}.").bold = True
                if item.get("area_score") is not None:
                    cp.add_run(f" ({item['area_score']}分)").font.size = Pt(9)
                
                widget = config.get("widget")
                
                if widget == "input_line":
                    lines = config.get("lines") or 1
                    # Set row height for input lines to ensure spacing
                    frame_table.rows[0].height = Cm(1.0 * lines)
                    frame_table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                    
                    # Add content (optional, just lines)
                    # For input_line, we might want actual lines inside?
                    # But the frame itself is the box.
                    # Let's add empty paragraphs with spacing
                    for _ in range(lines):
                         p = cell.add_paragraph()
                         p.paragraph_format.space_before = Pt(12)
                         p.paragraph_format.space_after = Pt(12)
                
                elif widget == "text_area":
                    rows = config.get("rows") or 5
                    # Set fixed height based on rows
                    # Approx 0.8cm per row
                    total_height = 0.8 * rows
                    frame_table.rows[0].height = Cm(total_height)
                    frame_table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                    
                    # Add empty paragraphs to push content if needed, but height rule handles frame size
                    # cell.add_paragraph() 
                        
                elif widget == "grid_area":
                    style_type = config.get("style")
                    if style_type == "line":
                         rows = config.get("rows") or 10
                         for _ in range(rows):
                             cp = cell.add_paragraph()
                             cp.paragraph_format.space_after = Pt(12)
                             cp.add_run("_" * 45).font.color.rgb = RGBColor(200, 0, 0)
                    else:
                        # Grid placeholder
                        cell.add_paragraph("[作文方格区域]")
                        frame_table.rows[0].height = Cm(5) # Default height for grid
                        
                doc.add_paragraph() # Spacer between frames
            except Exception as e:
                 print(f"Error rendering sub item {item.get('question_id')}: {e}")
                 doc.add_paragraph(f"[题目渲染错误: {item.get('area_sort')}]")

    return doc


@answer_sheets_bp.post("/<int:sheet_id>/export/word")
def export_sheet_word(sheet_id: int):
    payload = request.get_json(silent=True) or {}
    paper_size = payload.get("paper_size", "A3") # Default A3
    ticket_no_digits = int(payload.get("ticket_no_digits", 10))

    t_sheet = _table("exam_answer_sheet")
    t_rel = _table("sheet_question_relation")
    t_style = _table("answer_area_style")
    session = get_session(current_app)
    
    try:
        sheet = session.execute(select(t_sheet).where(t_sheet.c.sheet_id == sheet_id)).mappings().first()
        if not sheet:
            return jsonify({"error": {"message": "答题卡不存在", "type": "NotFound"}}), 404
            
        items = session.execute(select(t_rel).where(t_rel.c.sheet_id == sheet_id).order_by(t_rel.c.area_sort)).mappings().all()
        items = [dict(r) for r in items]
        
        # Fetch styles
        style_ids = set(it["style_id"] for it in items if it["style_id"])
        styles = {}
        if style_ids:
            s_rows = session.execute(select(t_style).where(t_style.c.style_id.in_(style_ids))).mappings().all()
            styles = {r["style_id"]: dict(r) for r in s_rows}
            
        doc = _render_sheet_word(dict(sheet), items, styles, paper_size=paper_size, ticket_no_digits=ticket_no_digits)
        
        # Stream directly
        f = BytesIO()
        doc.save(f)
        f.seek(0)
        
        filename = f"{sheet.get('sheet_name', '答题卡')}.docx"
        try:
            from urllib.parse import quote
            filename = quote(filename)
        except:
            pass
            
        return send_file(f, as_attachment=True, download_name=f"{sheet.get('sheet_name', '答题卡')}.docx")
        
    except SQLAlchemyError as err:
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500
    except Exception as e:
        traceback.print_exc() # Log full trace
        return jsonify({"error": {"message": f"Word生成失败: {str(e)}", "type": "GenerationError"}}), 500


@answer_sheets_bp.post("/<int:sheet_id>/export/pdf")
def export_sheet_pdf(sheet_id: int):
    if docx2pdf_convert is None:
        return jsonify({"error": {"message": "PDF 导出不可用：docx2pdf 未安装", "type": "NotSupported"}}), 400
        
    payload = request.get_json(silent=True) or {}
    paper_size = payload.get("paper_size", "A3")
    ticket_no_digits = int(payload.get("ticket_no_digits", 10))

    t_sheet = _table("exam_answer_sheet")
    t_rel = _table("sheet_question_relation")
    t_style = _table("answer_area_style")
    session = get_session(current_app)
    
    try:
        sheet = session.execute(select(t_sheet).where(t_sheet.c.sheet_id == sheet_id)).mappings().first()
        if not sheet:
            return jsonify({"error": {"message": "答题卡不存在", "type": "NotFound"}}), 404
            
        items = session.execute(select(t_rel).where(t_rel.c.sheet_id == sheet_id).order_by(t_rel.c.area_sort)).mappings().all()
        items = [dict(r) for r in items]
        
        style_ids = set(it["style_id"] for it in items if it["style_id"])
        styles = {}
        if style_ids:
            s_rows = session.execute(select(t_style).where(t_style.c.style_id.in_(style_ids))).mappings().all()
            styles = {r["style_id"]: dict(r) for r in s_rows}
            
        doc = _render_sheet_word(dict(sheet), items, styles, paper_size=paper_size, ticket_no_digits=ticket_no_digits)
        
        # Use temp files for conversion
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
            doc.save(tmp_docx.name)
            tmp_docx_path = tmp_docx.name
            
        tmp_pdf_path = tmp_docx_path.replace(".docx", ".pdf")
        
        try:
            docx2pdf_convert(tmp_docx_path, tmp_pdf_path)
            
            with open(tmp_pdf_path, "rb") as f:
                pdf_data = f.read()
                
            return send_file(
                BytesIO(pdf_data), 
                as_attachment=True, 
                download_name=f"{sheet.get('sheet_name', '答题卡')}.pdf",
                mimetype='application/pdf'
            )
        finally:
            if os.path.exists(tmp_docx_path):
                os.remove(tmp_docx_path)
            if os.path.exists(tmp_pdf_path):
                os.remove(tmp_pdf_path)
                
    except Exception as e:
        return jsonify({"error": {"message": f"PDF转换失败: {str(e)}", "type": "ConversionError"}}), 500


@answer_sheets_bp.get("/download/<int:sheet_id>/<filename>")
def download_file(sheet_id: int, filename: str):
    # Security check: ensure filename belongs to sheet_id dir
    base_dir = _export_base_dir(f"sheet_{sheet_id}")
    path = os.path.join(base_dir, filename)
    if os.path.exists(path):
        return send_file(path, as_attachment=True)
    return jsonify({"error": {"message": "文件不存在", "type": "NotFound"}}), 404



@answer_styles_bp.get("")
def list_styles():
    type_id = request.args.get("type_id", type=int)
    t = _table("answer_area_style")
    stmt = select(t.c.style_id, t.c.type_id, t.c.style_name, t.c.style_config, t.c.is_default, t.c.create_time)
    if type_id is not None:
        stmt = stmt.where(t.c.type_id == type_id)
    stmt = stmt.order_by(t.c.type_id.asc(), t.c.is_default.desc(), t.c.style_id.desc())
    try:
        rows = get_session(current_app).execute(stmt).mappings().all()
        return jsonify({"items": [dict(r) for r in rows]})
    except SQLAlchemyError as err:
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@answer_styles_bp.post("")
def create_style():
    payload = request.get_json(silent=True) or {}
    t = _table("answer_area_style")
    if not payload.get("type_id") or not payload.get("style_name"):
        return jsonify({"error": {"message": "type_id 与 style_name 必填", "type": "BadRequest"}}), 400

    data = {
        "type_id": int(payload["type_id"]),
        "style_name": payload["style_name"],
        "style_config": payload.get("style_config") or "{}",
        "is_default": 1 if payload.get("is_default") else 0,
        "create_time": datetime.now(),
    }

    session = get_session(current_app)
    try:
        if data["is_default"] == 1:
            session.execute(update(t).where(t.c.type_id == data["type_id"]).values(is_default=0))
        res = session.execute(insert(t).values(**data))
        session.commit()
        style_id = res.inserted_primary_key[0] if res.inserted_primary_key else None
        return jsonify({"style_id": style_id})
    except IntegrityError as err:
        session.rollback()
        return jsonify({"error": {"message": str(err), "type": "IntegrityError"}}), 400
    except SQLAlchemyError as err:
        session.rollback()
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@answer_styles_bp.put("/<int:style_id>")
def update_style(style_id: int):
    payload = request.get_json(silent=True) or {}
    t = _table("answer_area_style")
    data = {k: payload.get(k) for k in ["type_id", "style_name", "style_config"] if k in payload}
    if "is_default" in payload:
        data["is_default"] = 1 if payload.get("is_default") else 0
    if not data:
        return jsonify({"ok": True})

    session = get_session(current_app)
    try:
        if data.get("is_default") == 1:
            row = session.execute(select(t.c.type_id).where(t.c.style_id == style_id)).scalar_one_or_none()
            if row is not None:
                session.execute(update(t).where(t.c.type_id == row).values(is_default=0))
        session.execute(update(t).where(t.c.style_id == style_id).values(**data))
        session.commit()
        return jsonify({"ok": True})
    except IntegrityError as err:
        session.rollback()
        return jsonify({"error": {"message": str(err), "type": "IntegrityError"}}), 400
    except SQLAlchemyError as err:
        session.rollback()
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@answer_styles_bp.delete("/<int:style_id>")
def delete_style(style_id: int):
    t = _table("answer_area_style")
    session = get_session(current_app)
    try:
        session.execute(delete(t).where(t.c.style_id == style_id))
        session.commit()
        return jsonify({"ok": True})
    except SQLAlchemyError as err:
        session.rollback()
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@answer_sheets_bp.get("")
def list_sheets():
    paper_id = request.args.get("paper_id", type=int)
    t = _table("exam_answer_sheet")
    stmt = select(t.c.sheet_id, t.c.paper_id, t.c.sheet_name, t.c.template_config, t.c.create_user, t.c.create_time, t.c.update_time)
    if paper_id is not None:
        stmt = stmt.where(t.c.paper_id == paper_id)
    stmt = stmt.order_by(t.c.sheet_id.desc())
    try:
        rows = get_session(current_app).execute(stmt).mappings().all()
        return jsonify({"items": [dict(r) for r in rows]})
    except SQLAlchemyError as err:
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@answer_sheets_bp.post("")
def create_sheet():
    payload = request.get_json(silent=True) or {}
    t = _table("exam_answer_sheet")
    if not payload.get("paper_id") or not payload.get("sheet_name"):
        return jsonify({"error": {"message": "paper_id 与 sheet_name 必填", "type": "BadRequest"}}), 400

    data = {
        "paper_id": int(payload["paper_id"]),
        "sheet_name": payload["sheet_name"],
        "template_config": payload.get("template_config") or "{}",
        "create_user": payload.get("create_user") or "creator",
        "create_time": datetime.now(),
        "update_time": datetime.now(),
    }
    session = get_session(current_app)
    try:
        res = session.execute(insert(t).values(**data))
        session.commit()
        sheet_id = res.inserted_primary_key[0] if res.inserted_primary_key else None
        return jsonify({"sheet_id": sheet_id})
    except IntegrityError as err:
        session.rollback()
        return jsonify({"error": {"message": str(err), "type": "IntegrityError"}}), 400
    except SQLAlchemyError as err:
        session.rollback()
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@answer_sheets_bp.put("/<int:sheet_id>")
def update_sheet(sheet_id: int):
    payload = request.get_json(silent=True) or {}
    t = _table("exam_answer_sheet")
    data = {k: payload.get(k) for k in ["sheet_name", "template_config", "create_user"] if k in payload}
    if not data:
        return jsonify({"ok": True})
    data["update_time"] = datetime.now()
    session = get_session(current_app)
    try:
        session.execute(update(t).where(t.c.sheet_id == sheet_id).values(**data))
        session.commit()
        return jsonify({"ok": True})
    except SQLAlchemyError as err:
        session.rollback()
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@answer_sheets_bp.delete("/<int:sheet_id>")
def delete_sheet(sheet_id: int):
    t = _table("exam_answer_sheet")
    rel = _table("sheet_question_relation")
    session = get_session(current_app)
    try:
        session.execute(delete(rel).where(rel.c.sheet_id == sheet_id))
        session.execute(delete(t).where(t.c.sheet_id == sheet_id))
        session.commit()
        return jsonify({"ok": True})
    except SQLAlchemyError as err:
        session.rollback()
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@answer_sheets_bp.get("/<int:sheet_id>/items")
def list_sheet_items(sheet_id: int):
    rel = _table("sheet_question_relation")
    stmt = select(
        rel.c.relation_id,
        rel.c.sheet_id,
        rel.c.question_id,
        rel.c.style_id,
        rel.c.area_sort,
        rel.c.area_score,
        rel.c.create_time,
    ).where(rel.c.sheet_id == sheet_id).order_by(rel.c.area_sort.asc())
    try:
        rows = get_session(current_app).execute(stmt).mappings().all()
        return jsonify({"items": [dict(r) for r in rows]})
    except SQLAlchemyError as err:
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@answer_sheets_bp.put("/<int:sheet_id>/items")
def update_sheet_items(sheet_id: int):
    payload = request.get_json(silent=True) or {}
    items = payload.get("items") or []
    if not isinstance(items, list):
        return jsonify({"error": {"message": "items 必须是数组", "type": "BadRequest"}}), 400

    rel = _table("sheet_question_relation")
    session = get_session(current_app)
    try:
        for it in items:
            if not it.get("question_id"):
                continue
            data = {}
            for k in ["style_id", "area_sort", "area_score"]:
                if k in it:
                    data[k] = it.get(k)
            if data:
                session.execute(
                    update(rel)
                    .where(and_(rel.c.sheet_id == sheet_id, rel.c.question_id == int(it["question_id"])))
                    .values(**data)
                )
        session.commit()
        return jsonify({"ok": True})
    except SQLAlchemyError as err:
        session.rollback()
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@answer_sheets_bp.post("/from-paper/<int:paper_id>")
def create_sheet_from_paper(paper_id: int):
    payload = request.get_json(silent=True) or {}
    create_user = payload.get("create_user") or "creator"
    template_config = payload.get("template_config") or "{}"

    paper = _table("exam_paper")
    pqr = _table("paper_question_relation")
    qb = _table("question_bank")
    sheet = _table("exam_answer_sheet")
    rel = _table("sheet_question_relation")
    style = _table("answer_area_style")

    session = get_session(current_app)
    now = datetime.now()

    try:
        p = session.execute(select(paper).where(paper.c.paper_id == paper_id)).mappings().first()
        if p is None:
            return jsonify({"error": {"message": "试卷不存在", "type": "NotFound"}}), 404

        existing = session.execute(select(sheet).where(sheet.c.paper_id == paper_id)).mappings().first()
        if existing:
            sheet_id = existing["sheet_id"]
            session.execute(delete(rel).where(rel.c.sheet_id == sheet_id))
            session.execute(update(sheet).where(sheet.c.sheet_id == sheet_id).values(update_time=now))
        else:
            res = session.execute(
                insert(sheet).values(
                    paper_id=paper_id,
                    sheet_name=f"{p.get('paper_name')}-答题卡",
                    template_config=template_config,
                    create_user=create_user,
                    create_time=now,
                    update_time=now,
                )
            )
            sheet_id = res.inserted_primary_key[0] if res.inserted_primary_key else None
            if not sheet_id:
                raise RuntimeError("创建答题卡失败")

        q_stmt = (
            select(pqr.c.question_id, pqr.c.question_sort, pqr.c.question_score, qb.c.type_id)
            .join(qb, qb.c.question_id == pqr.c.question_id)
            .where(pqr.c.paper_id == paper_id)
            .order_by(pqr.c.question_sort.asc())
        )
        qs = session.execute(q_stmt).mappings().all()

        default_styles = {
            int(r["type_id"]): int(r["style_id"])
            for r in session.execute(select(style.c.type_id, style.c.style_id).where(style.c.is_default == 1)).mappings().all()
        }

        for q in qs:
            type_id = int(q["type_id"])
            session.execute(
                insert(rel).values(
                    sheet_id=sheet_id,
                    question_id=int(q["question_id"]),
                    style_id=default_styles.get(type_id),
                    area_sort=int(q["question_sort"]),
                    area_score=q["question_score"],
                    create_time=now,
                )
            )

        session.commit()
        return jsonify({"sheet_id": sheet_id, "item_count": len(qs)})
    except IntegrityError as err:
        session.rollback()
        return jsonify({"error": {"message": str(err), "type": "IntegrityError"}}), 400
    except SQLAlchemyError as err:
        session.rollback()
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500
