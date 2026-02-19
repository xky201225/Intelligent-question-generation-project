from __future__ import annotations

import json
import os
import re
import threading
import time
import uuid
import io
import pdfplumber
from docx import Document
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime
from typing import Optional

from flask import Blueprint, Response, current_app, jsonify, request, stream_with_context
from sqlalchemy import and_, insert, select, update, func, or_
from sqlalchemy.exc import SQLAlchemyError

from app.db import get_db, get_session
from app.services.deepseek import get_deepseek_client

ai_bp = Blueprint("ai", __name__)

_executor = ThreadPoolExecutor(max_workers=2)
_jobs_lock = threading.Lock()
_jobs: dict[str, dict] = {}

MAX_FILL_ATTEMPTS = 3


def _table(name: str):
    db = get_db(current_app)
    if name not in db.metadata.tables:
        db.metadata.reflect(bind=db.engine, only=[name])
    return db.metadata.tables[name]


def _extract_json_list(text: str) -> list[dict]:
    text = (text or "").strip()

    # 去掉常见 markdown code fence
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z0-9_-]*\s*", "", text).strip()
        text = re.sub(r"\s*```$", "", text).strip()

    start = text.find("[")
    if start < 0:
        raise ValueError("AI 返回内容不是 JSON 数组：未找到 [")

    def find_balanced_end(s: str, start_idx: int) -> Optional[int]:
        in_string = False
        escape = False
        depth = 0
        for i in range(start_idx, len(s)):
            ch = s[i]
            if in_string:
                if escape:
                    escape = False
                    continue
                if ch == "\\":
                    escape = True
                    continue
                if ch == '"':
                    in_string = False
                continue
            if ch == '"':
                in_string = True
                continue
            if ch == "[":
                depth += 1
            elif ch == "]":
                depth -= 1
                if depth == 0:
                    return i
        return None

    end = find_balanced_end(text, start)
    candidate = text[start : end + 1] if end is not None else text[start:]

    try:
        parsed = json.loads(candidate)
        if isinstance(parsed, list):
            return parsed
    except Exception:
        pass

    # 尝试“容错解析”：逐个 raw_decode，尽量取到前面完整的对象
    decoder = json.JSONDecoder()
    items: list[dict] = []
    idx = 1  # 跳过 '['
    while True:
        while idx < len(candidate) and candidate[idx] in " \t\r\n,":
            idx += 1
        if idx >= len(candidate):
            break
        if candidate[idx] == "]":
            break
        try:
            obj, next_idx = decoder.raw_decode(candidate, idx)
        except Exception:
            break
        if isinstance(obj, dict):
            items.append(obj)
        idx = next_idx
    if items:
        return items

    raise ValueError("AI 返回内容不是有效 JSON 数组")


def _job_update(job_id: str, patch: dict) -> None:
    with _jobs_lock:
        if job_id not in _jobs:
            return
        _jobs[job_id].update(patch)


def _job_snapshot(job_id: str) -> Optional[dict]:
    with _jobs_lock:
        v = _jobs.get(job_id)
        return dict(v) if v else None


def _job_event(job_id: str, event_type: str, message: Optional[str] = None, data: Optional[dict] = None) -> None:
    now = datetime.now().isoformat(timespec="seconds")
    with _jobs_lock:
        job = _jobs.get(job_id)
        if not job:
            return
        seq = int(job.get("seq") or 0) + 1
        job["seq"] = seq
        ev = {"id": seq, "ts": now, "type": event_type, "message": message, "data": data or {}}
        job.setdefault("events", []).append(ev)
        if len(job["events"]) > 4000:
            job["events"] = job["events"][-2000:]


def _pick_first(d: dict, keys: list[str]) -> Optional[str]:
    for k in keys:
        v = d.get(k)
        if v is None:
            continue
        s = str(v).strip()
        if s != "":
            return s
    return None


def _run_generation(app, job_id: str, subject_id: int, chapter_ids: list[int], rules: list[dict], create_user: str):
    with app.app_context():
        _job_update(job_id, {"status": "running", "started_at": datetime.now().isoformat(timespec="seconds")})
        _job_event(job_id, "job_start", "开始生成")
        qb = _table("question_bank")
        ch = _table("textbook_chapter")
        session = get_session(current_app)
        inserted = 0
        created_ids: list[int] = []

        try:
            stmt = select(ch.c.chapter_id, ch.c.chapter_name, ch.c.content).where(ch.c.chapter_id.in_(chapter_ids))
            chapters = session.execute(stmt).mappings().all()
            if not chapters:
                _job_update(job_id, {"status": "error", "error": "所选章节不存在"})
                _job_event(job_id, "job_error", "所选章节不存在")
                return

            for chapter in chapters:
                chapter_id = int(chapter["chapter_id"])
                summary = chapter.get("content") or ""
                _job_event(job_id, "chapter_start", f"章节：{chapter.get('chapter_name')}", {"chapter_id": chapter_id})

                for rule in rules:
                    type_id = rule["type_id"]
                    difficulty_id = rule["difficulty_id"]
                    count = rule["count"]
                    _job_event(
                        job_id,
                        "rule_start",
                        f"题型={type_id} 难度={difficulty_id} 数量={count}",
                        {"chapter_id": chapter_id, "type_id": type_id, "difficulty_id": difficulty_id, "count": count},
                    )

                    sample_stmt = (
                        select(qb.c.question_id, qb.c.question_content, qb.c.question_answer, qb.c.question_analysis)
                        .where(and_(qb.c.chapter_id == chapter_id, qb.c.review_status == 1, qb.c.type_id == type_id))
                        .order_by(qb.c.question_id.desc())
                        .limit(3)
                    )
                    sample = session.execute(sample_stmt).mappings().all()
                    source_ids = [str(s["question_id"]) for s in sample]

                    system_prompt = "你是一位大学出题助理。你必须严格输出 JSON 数组，不要输出任何多余文本，严格按照题型输出，不要出错题型，除了单选题和多项题任何题型都不要有选项。"
                    
                    additional_reqs = ""
                    if type_id == 1: # 单选题
                        additional_reqs = "5) 必须有且只能有4个选项(A/B/C/D)。\n6) 必须且只能有一个正确答案。\n7) 【强制】题干末尾必须以中文括号（ ）结尾。\n"
                    elif type_id == 7: # 多选题
                        additional_reqs = "5) 必须有且只能有4个选项(A/B/C/D)。\n6) 必须有两个或更多正确答案。\n7) 【强制】题干末尾必须以中文括号（ ）结尾。\n"
                    elif type_id == 2: # 判断题
                        additional_reqs = "5) 题干必须是陈述句。\n6) 答案必须是“正确”或“错误”（或T/F）。\n7) 【强制】题干末尾必须以中文括号（ ）结尾。\n"
                    elif type_id == 4: # 计算题
                        additional_reqs = "5) 必须是计算类题目，禁止出现选项(A/B/C/D)。\n6) 答案需包含具体计算结果。\n"
                    elif type_id == 3: # 填空题
                        additional_reqs = "5) 题干中必须包含填空符（如______）。\n6) 禁止出现选项。\n"
                    elif type_id in [5, 6]: # 简答题, 作文
                        additional_reqs = "5) 必须是主观题，禁止出现选项。\n"

                    content_req = "4) question_content 内可包含选项（如A/B/C/D），但仍是纯文本\n"
                    if type_id not in [1, 7]:
                        content_req = "4) question_content 必须是纯文本，严禁包含选项(A/B/C/D)\n"

                    user_prompt = (
                        "请为指定教材章节生成题目，要求：\n"
                        "1) 只生成与章节相关的题\n"
                        "2) 难度与题型遵循要求\n"
                        "3) 输出严格 JSON 数组，每个元素包含字段：type_id, question_content, question_answer, question_analysis, question_score\n"
                        f"{content_req}"
                        f"{additional_reqs}\n"
                        f"章节名称：{chapter['chapter_name']}\n"
                        f"章节概要：{summary if summary else '(暂无概要)'}\n"
                        f"目标数量：{count}\n"
                        f"题型ID：{type_id} (请在返回的JSON中将 type_id 设为 {type_id})\n"
                        f"难度ID：{difficulty_id}\n\n"
                        "参考题目（用于风格与覆盖点，不要重复）：\n"
                    )
                    for i, s in enumerate(sample, start=1):
                        user_prompt += (
                            f"\n{i}. 题干：{s.get('question_content')}\n"
                            f"   答案：{s.get('question_answer')}\n"
                            f"   解析：{s.get('question_analysis')}\n"
                        )

                    client = get_deepseek_client()
                    target_count = int(count)
                    collected: list[dict] = []
                    seen_contents: set[str] = set()
                    attempt = 0

                    def call_model(prompt: str, attempt_no: int, missing: int) -> str:
                        _job_event(
                            job_id,
                            "ai_start",
                            f"请求模型中…（第{attempt_no}次，缺{missing}题）",
                            {"chapter_id": chapter_id, "type_id": type_id, "difficulty_id": difficulty_id, "attempt": attempt_no, "missing": missing},
                        )
                        try:
                            raw_chunks: list[str] = []
                            buf = ""
                            last_flush = time.time()
                            for chunk in client.chat_stream(system_prompt=system_prompt, user_prompt=prompt, temperature=0.7):
                                raw_chunks.append(chunk)
                                buf += chunk
                                now_ts = time.time()
                                if len(buf) >= 200 or "\n" in buf or (now_ts - last_flush) >= 0.8:
                                    _job_event(job_id, "ai_delta", data={"text": buf})
                                    buf = ""
                                    last_flush = now_ts
                            if buf:
                                _job_event(job_id, "ai_delta", data={"text": buf})
                            raw_text = "".join(raw_chunks)
                        except Exception as err:
                            _job_event(job_id, "ai_error", f"流式输出不可用，改用普通请求：{err}")
                            raw_text = client.chat(system_prompt=system_prompt, user_prompt=prompt, temperature=0.7)
                            _job_event(job_id, "ai_delta", data={"text": raw_text})
                        _job_event(job_id, "ai_end", "模型返回完成")
                        return raw_text

                    while len(collected) < target_count and attempt < MAX_FILL_ATTEMPTS:
                        attempt += 1
                        missing = target_count - len(collected)

                        prompt = user_prompt.replace(f"目标数量：{count}\n", f"目标数量：{missing}\n")
                        if attempt > 1:
                            existed = "\n".join([f"- {x['question_content']}" for x in collected[:50]])
                            prompt += (
                                "\n补齐要求：\n"
                                f"1) 你只需要补齐缺口：再生成 {missing} 道新题\n"
                                "2) 只输出 JSON 数组，不要输出 ```json 或任何解释文本\n"
                                "3) 绝对不要重复已有题干\n"
                                f"\n已生成题干（不要重复）：\n{existed}\n"
                            )

                        _job_event(job_id, "rule_retry", f"补齐生成：第{attempt}次（还差{missing}题）")
                        raw = call_model(prompt, attempt, missing)
                        try:
                            items = _extract_json_list(raw)
                        except Exception as err:
                            _job_event(job_id, "rule_warn", f"解析失败（第{attempt}次）：{err}")
                            continue

                        _job_event(job_id, "parse_ok", f"解析成功：{len(items)}条（第{attempt}次）", {"count": len(items), "attempt": attempt})

                        for it in items:
                            if len(collected) >= target_count:
                                break
                            if not isinstance(it, dict):
                                continue
                            content = _pick_first(it, ["question_content", "content", "stem", "question", "题干"])
                            if not content:
                                continue
                            content_key = re.sub(r"\s+", " ", content).strip()
                            if content_key in seen_contents:
                                continue
                            seen_contents.add(content_key)

                            answer = _pick_first(it, ["question_answer", "answer", "答案"])
                            analysis = _pick_first(it, ["question_analysis", "analysis", "解析"])
                            score_raw = _pick_first(it, ["question_score", "score", "分值"])
                            score_val = None
                            if score_raw is not None:
                                try:
                                    score_val = float(score_raw)
                                except Exception:
                                    score_val = None

                            returned_type_id = _pick_first(it, ["type_id", "type", "题型ID"])
                            final_type_id = type_id
                            if returned_type_id:
                                try:
                                    final_type_id = int(returned_type_id)
                                except:
                                    pass

                            collected.append(
                                {
                                    "question_content": content,
                                    "question_answer": answer,
                                    "question_analysis": analysis,
                                    "question_score": score_val,
                                    "type_id": final_type_id,
                                }
                            )

                        _job_event(job_id, "rule_progress", f"已收集：{len(collected)}/{target_count}", {"collected": len(collected), "target": target_count})

                    if len(collected) < target_count:
                        _job_event(job_id, "rule_error", f"补齐失败：需要{target_count}，实际{len(collected)}（已重试{attempt}次）")
                        raise ValueError(f"生成题目不足：需要{target_count}，实际{len(collected)}")

                    now = datetime.now()
                    for it in collected:
                        data = {
                            "subject_id": subject_id,
                            "chapter_id": chapter_id,
                            "type_id": it.get("type_id", type_id),
                            "difficulty_id": difficulty_id,
                            "question_content": it.get("question_content"),
                            "question_answer": it.get("question_answer"),
                            "question_analysis": it.get("question_analysis"),
                            "question_score": it.get("question_score") if it.get("question_score") is not None else 0,
                            "is_ai_generated": 1,
                            "source_question_ids": ",".join(source_ids) if source_ids else None,
                            "review_status": 0,
                            "reviewer": None,
                            "review_time": None,
                            "create_user": create_user,
                            "create_time": now,
                            "update_time": now,
                        }
                        res = session.execute(insert(qb).values(**data))
                        if res.inserted_primary_key:
                            created_ids.append(res.inserted_primary_key[0])
                        inserted += 1

                    _job_update(job_id, {"inserted": inserted, "question_ids": created_ids})
                    _job_event(job_id, "progress", f"已入库：{inserted}题", {"inserted": inserted})
                    _job_event(job_id, "rule_end", "本规则完成")

            session.commit()
            _job_update(
                job_id,
                {
                    "status": "done",
                    "inserted": inserted,
                    "question_ids": created_ids,
                    "finished_at": datetime.now().isoformat(timespec="seconds"),
                },
            )
            _job_event(job_id, "job_done", f"任务完成：新增{inserted}题", {"inserted": inserted})
        except Exception as err:
            session.rollback()
            _job_update(
                job_id,
                {
                    "status": "error",
                    "error": str(err),
                    "inserted": inserted,
                    "question_ids": created_ids,
                    "finished_at": datetime.now().isoformat(timespec="seconds"),
                },
            )
            _job_event(job_id, "job_error", str(err))


def _extract_file_content(file) -> str:
    """Extract text content from uploaded file (PDF/Word)"""
    filename = file.filename.lower()
    content = ""
    
    try:
        if filename.endswith(".pdf"):
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        content += text + "\n"
        elif filename.endswith(".docx") or filename.endswith(".doc"):
            doc = Document(file)
            for para in doc.paragraphs:
                content += para.text + "\n"
        else:
            # Try text decode
            content = file.read().decode("utf-8", errors="ignore")
    except Exception as e:
        raise ValueError(f"文件解析失败: {str(e)}")
        
    return content.strip()


def _run_variant_generation(app, job_id: str, subject_id: int, chapter_ids: list[int], source_content: str, create_user: str, source_type: str = "text", predefined_tasks: list[dict] = None):
    """
    Variant generation worker
    source_content: Raw text (from file) or formatted questions (from paper)
    predefined_tasks: Optional list of tasks [{"chapter_id": 1, "chapter_name": "...", "source_text": "...", "target_count": 5}]
    """
    with app.app_context():
        _job_update(job_id, {"status": "running", "started_at": datetime.now().isoformat(timespec="seconds")})
        _job_event(job_id, "job_start", "开始生成变式题目")
        
        qb = _table("question_bank")
        ch = _table("textbook_chapter")
        session = get_session(current_app)
        inserted = 0
        created_ids: list[int] = []

        try:
            # Prepare tasks
            tasks = []
            if predefined_tasks:
                tasks = predefined_tasks
            elif chapter_ids:
                stmt = select(ch.c.chapter_id, ch.c.chapter_name, ch.c.content).where(ch.c.chapter_id.in_(chapter_ids))
                rows = session.execute(stmt).mappings().all()
                if not rows:
                    raise ValueError("所选章节不存在")
                # Default target: 5 per chapter
                for r in rows:
                    tasks.append({
                        "chapter_id": int(r["chapter_id"]),
                        "chapter_name": r["chapter_name"],
                        "chapter_summary": r.get("content") or "",
                        "target_count": 5,
                        "source_text": source_content # Use global source content if not predefined
                    })
            else:
                # No chapter selected: Global generation based on source
                count = 10
                # Try to estimate count from source content (works for both paper and file)
                matches = re.findall(r"^\d+[\.、\s]", source_content, re.MULTILINE)
                if matches:
                    count = len(matches)
                
                tasks.append({
                    "chapter_id": None,
                    "chapter_name": "通用（无特定章节）",
                    "chapter_summary": "",
                    "target_count": count,
                    "source_text": source_content
                })

            total_tasks = sum(t["target_count"] for t in tasks)
            _job_event(job_id, "job_start", f"开始生成变式题目（共{total_tasks}题）", {"total_count": total_tasks})

            for task in tasks:
                cid = task["chapter_id"]
                cname = task["chapter_name"]
                csummary = task["chapter_summary"]
                target_count_per_chapter = task["target_count"]
                current_source = task.get("source_text") or source_content
                
                _job_event(job_id, "chapter_start", f"正在生成变式（章节：{cname}）...", {"chapter_id": cid})

                system_prompt = (
                    "你是专业的试题变式生成助手。你的任务是分析给定的【源材料】（Source Material），"
                    "识别其中的题型、难度和出题风格，然后生成新的变式题目。\n"
                    "要求：\n"
                    "1. 严格输出 JSON 数组，不要包含 markdown 标记或额外文本。\n"
                    "2. 题型必须映射到系统ID：1=单选, 2=判断, 3=填空, 4=计算, 5=简答, 6=作文, 7=多选。单选和多选有且必须只有ABCD四个选项\n"
                    "3. 题目内容应尽可能模仿【源材料】的风格。\n"
                    "4. JSON字段：type_id, question_content, question_answer, question_analysis, question_score (可选).\n"
                    "5. 【强制格式】：判断题、单选题、多选题的题干末尾，必须以中文括号（ ）结尾，不得遗漏。\n"
                    "6. 【选项处理】：对于单选题(1)和多选题(7)，选项(A. B. C. D.)必须完整包含在 question_content 中，换行显示。\n"
                    "7.生成的题目数量必须严格等于要求的数量，且不得包含任何与要求不符的题型。"
                )

                user_prompt = (
                    f"【目标范围】\n名称：{cname}\n概要：{csummary[:500]}\n\n"
                    f"【源材料】\n{current_source[:3000]}...\n\n" # Limit context to avoid overflow
                    f"【任务】\n请基于源材料的风格，生成 {target_count_per_chapter} 道变式题目。\n"
                    "如果源材料包含多种题型，请尽量覆盖。\n"
                    "输出 JSON 数组："
                )

                client = get_deepseek_client()
                collected = []
                seen_contents = set()
                attempt = 0
                
                while len(collected) < target_count_per_chapter and attempt < 3:
                    attempt += 1
                    missing = target_count_per_chapter - len(collected)
                    _job_event(job_id, "ai_start", f"正在请求模型（第{attempt}次，缺{missing}题）...")
                    
                    # 动态调整 Prompt
                    prompt = user_prompt
                    if attempt > 1:
                         prompt += f"\n\n【重要】你只需要再生成 {missing} 道题即可，绝对不要重复已有题目。"

                    try:
                        raw_chunks = []
                        buf = ""
                        last_flush = time.time()
                        for chunk in client.chat_stream(system_prompt=system_prompt, user_prompt=prompt, temperature=0.7):
                            raw_chunks.append(chunk)
                            buf += chunk
                            now_ts = time.time()
                            if len(buf) >= 200 or "\n" in buf or (now_ts - last_flush) >= 0.8:
                                _job_event(job_id, "ai_delta", data={"text": buf})
                                buf = ""
                                last_flush = now_ts
                        if buf:
                            _job_event(job_id, "ai_delta", data={"text": buf})
                        raw_text = "".join(raw_chunks)
                        
                        items = _extract_json_list(raw_text)
                        
                        for it in items:
                            if len(collected) >= target_count_per_chapter: break
                            
                            # Validate and Normalize
                            content = _pick_first(it, ["question_content", "content", "题干"])
                            if not content: continue
                            
                            # Check duplicates
                            if content in seen_contents: continue
                            seen_contents.add(content)
                            
                            answer = _pick_first(it, ["question_answer", "answer", "答案"])
                            analysis = _pick_first(it, ["question_analysis", "analysis", "解析"])
                            
                            tid = it.get("type_id")
                            if not tid:
                                # Try to guess or default to 5 (Short Answer) if unknown
                                tid = 5 
                            
                            collected.append({
                                "question_content": content,
                                "question_answer": answer,
                                "question_analysis": analysis,
                                "type_id": int(tid),
                                "question_score": it.get("question_score")
                            })
                            
                        _job_event(job_id, "parse_ok", f"解析成功：{len(items)}条")
                        
                    except Exception as e:
                        _job_event(job_id, "ai_error", f"生成出错：{str(e)}")
                
                # Insert into DB
                # Safety check
                if len(collected) > target_count_per_chapter:
                    collected = collected[:target_count_per_chapter]

                now = datetime.now()
                for it in collected:
                    data = {
                        "subject_id": subject_id,
                        "chapter_id": cid, # Can be None
                        "type_id": it["type_id"],
                        "difficulty_id": 3, # Default to medium if not detected
                        "question_content": it["question_content"],
                        "question_answer": it["question_answer"],
                        "question_analysis": it["question_analysis"],
                        "question_score": it.get("question_score") if it.get("question_score") is not None else 0,
                        "is_ai_generated": 1,
                        "review_status": 0,
                        "create_user": create_user,
                        "create_time": now,
                        "update_time": now,
                    }
                    res = session.execute(insert(qb).values(**data))
                    if res.inserted_primary_key:
                        created_ids.append(res.inserted_primary_key[0])
                    inserted += 1
                
                session.commit()
                _job_event(job_id, "progress", f"章节/任务 {cname} 完成，入库 {len(collected)} 题")

            _job_update(job_id, {
                "status": "done", 
                "inserted": inserted, 
                "question_ids": created_ids,
                "finished_at": datetime.now().isoformat(timespec="seconds")
            })
            _job_event(job_id, "job_done", f"全部完成，共生成 {inserted} 题")

        except Exception as e:
            session.rollback()
            _job_update(job_id, {"status": "error", "error": str(e)})
            _job_event(job_id, "job_error", str(e))


@ai_bp.post("/generate-from-paper")
def generate_from_paper():
    payload = request.get_json(silent=True) or {}
    paper_id = payload.get("paper_id")
    subject_id = payload.get("subject_id")
    chapter_ids = payload.get("chapter_ids") or []
    create_user = payload.get("create_user") or "ai"

    if not paper_id or not subject_id:
        return jsonify({"error": {"message": "paper_id 和 subject_id 必填", "type": "BadRequest"}}), 400

    # Fetch paper content
    try:
        session = get_session(current_app)
        # Join with question bank to get content
        pqr = _table("paper_question_relation")
        qb = _table("question_bank")
        ch = _table("textbook_chapter")
        
        stmt = (
            select(
                qb.c.question_content, 
                qb.c.type_id, 
                qb.c.question_score,
                qb.c.chapter_id,
                ch.c.chapter_name
            )
            .join(pqr, pqr.c.question_id == qb.c.question_id)
            .outerjoin(ch, ch.c.chapter_id == qb.c.chapter_id)
            .where(pqr.c.paper_id == paper_id)
            .order_by(pqr.c.question_sort)
        )
        rows = session.execute(stmt).mappings().all()
        
        if not rows:
            return jsonify({"error": {"message": "试卷为空或不存在", "type": "NotFound"}}), 404
            
        # Group by chapter_id
        # { chapter_id: { "name": str, "questions": [] } }
        grouped = {}
        for r in rows:
            cid = r["chapter_id"]
            if cid is None:
                cid = 0 # Use 0 for unknown/null chapter
            
            if cid not in grouped:
                grouped[cid] = {
                    "chapter_id": cid if cid != 0 else None,
                    "chapter_name": r["chapter_name"] if r["chapter_name"] else "未分类章节",
                    "questions": []
                }
            
            grouped[cid]["questions"].append(r)
        
        # Construct tasks
        tasks = []
        for cid, group in grouped.items():
            # Construct source text for this chapter
            source_text = f"【{group['chapter_name']} 题目列表】\n"
            for i, q in enumerate(group["questions"], 1):
                source_text += f"{i}. [类型ID:{q['type_id']}] {q['question_content']} ({q['question_score']}分)\n"
            
            tasks.append({
                "chapter_id": group["chapter_id"],
                "chapter_name": group["chapter_name"],
                "chapter_summary": "", # No summary available from this join, maybe fine
                "target_count": len(group["questions"]), # Generate same amount as source
                "source_text": source_text
            })
            
    except SQLAlchemyError as e:
        return jsonify({"error": {"message": str(e), "type": "DatabaseError"}}), 500

    job_id = uuid.uuid4().hex
    with _jobs_lock:
        _jobs[job_id] = {
            "job_id": job_id, "status": "queued", "inserted": 0, "question_ids": [],
            "created_at": datetime.now().isoformat(timespec="seconds"), "seq": 0, "events": []
        }

    app = current_app._get_current_object()
    _executor.submit(_run_variant_generation, app, job_id, int(subject_id), [], "", create_user, "paper", tasks)
    
    return jsonify({"ok": True, "job_id": job_id})


@ai_bp.post("/generate-from-file")
def generate_from_file():
    if "file" not in request.files:
        return jsonify({"error": {"message": "未上传文件", "type": "BadRequest"}}), 400
        
    file = request.files["file"]
    
    # Check size (10MB limit)
    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)
    if size > 10 * 1024 * 1024:
         return jsonify({"error": {"message": "文件大小不能超过 10MB", "type": "BadRequest"}}), 400

    filename = file.filename.lower() if file.filename else ""
    if not filename.endswith(('.pdf', '.doc', '.docx')):
         return jsonify({"error": {"message": "不支持的文件格式，仅支持 PDF/Word", "type": "BadRequest"}}), 400

    subject_id = request.form.get("subject_id")
    chapter_ids_str = request.form.get("chapter_ids")
    create_user = request.form.get("create_user") or "ai"
    
    if not subject_id:
        return jsonify({"error": {"message": "subject_id 必填", "type": "BadRequest"}}), 400
        
    try:
        if chapter_ids_str:
            chapter_ids = [int(x) for x in chapter_ids_str.split(",") if x.strip()]
        else:
            chapter_ids = []
    except:
        return jsonify({"error": {"message": "chapter_ids 格式错误", "type": "BadRequest"}}), 400

    try:
        content = _extract_file_content(file)
        if not content:
             return jsonify({"error": {"message": "文件内容为空或无法解析", "type": "BadRequest"}}), 400
    except Exception as e:
        return jsonify({"error": {"message": str(e), "type": "FileError"}}), 400

    job_id = uuid.uuid4().hex
    with _jobs_lock:
        _jobs[job_id] = {
            "job_id": job_id, "status": "queued", "inserted": 0, "question_ids": [],
            "created_at": datetime.now().isoformat(timespec="seconds"), "seq": 0, "events": []
        }

    app = current_app._get_current_object()
    _executor.submit(_run_variant_generation, app, job_id, int(subject_id), chapter_ids, content, create_user, "file")
    
    return jsonify({"ok": True, "job_id": job_id})


@ai_bp.post("/generate-questions")
def generate_questions():
    payload = request.get_json(silent=True) or {}
    subject_id = payload.get("subject_id")
    chapter_ids = payload.get("chapter_ids") or []  # 被选中的小章节ID
    selected_main_chapters = payload.get("selected_main_chapters") or []  # 选中的大章节ID
    if payload.get("chapter_id"):  # 兼容旧参数
        chapter_ids.append(payload.get("chapter_id"))
    create_user = payload.get("create_user") or "ai"
    rules = payload.get("rules")
    chapter_weights = payload.get("chapter_weights") or {}

    if not rules:
        type_id = payload.get("type_id")
        difficulty_id = payload.get("difficulty_id")
        count = payload.get("count")
        if type_id and difficulty_id and count:
            rules = [{"type_id": type_id, "difficulty_id": difficulty_id, "count": count}]
        else:
            return jsonify({"error": {"message": "缺少生成规则 (rules 或 type_id/difficulty_id/count)", "type": "BadRequest"}}), 400

    if not subject_id or not chapter_ids:
        return jsonify({"error": {"message": "subject_id 和 chapter_ids 必填", "type": "BadRequest"}}), 400

    validated_rules = []
    total_count_needed = 0
    for r in rules:
        tid = r.get("type_id")
        did = r.get("difficulty_id")
        cnt = r.get("count")
        if not tid or not did or not cnt:
            continue
        try:
            cnt = int(cnt)
            if cnt < 1 or cnt > 50:
                return jsonify({"error": {"message": f"数量范围 1-50 (type_id={tid})", "type": "BadRequest"}}), 400
            validated_rules.append({"type_id": tid, "difficulty_id": did, "count": cnt})
            total_count_needed += cnt
        except:
            return jsonify({"error": {"message": "count 格式错误", "type": "BadRequest"}}), 400
    if not validated_rules:
        return jsonify({"error": {"message": "有效生成规则为空", "type": "BadRequest"}}), 400

    # 新分配逻辑：先大章节平均，再大章节内小章节平分
    final_chapter_dist = {}  # {chapter_id: ratio}
    if selected_main_chapters and chapter_ids:
        main_count = len(selected_main_chapters)
        if main_count == 0:
            return jsonify({"error": {"message": "未选择大章节", "type": "BadRequest"}}), 400
        main_ratio = 1.0 / main_count
        # 查询所有小章节的父章节
        app = current_app._get_current_object()
        session = get_session(app)
        ch = _table("textbook_chapter")
        # 获取小章节的parent_chapter_id
        stmt = select(ch.c.chapter_id, ch.c.parent_chapter_id).where(ch.c.chapter_id.in_(chapter_ids))
        rows = session.execute(stmt).mappings().all()
        # 按大章节分组
        group = {}
        for row in rows:
            pid = row["parent_chapter_id"]
            cid = row["chapter_id"]
            if pid in selected_main_chapters:
                group.setdefault(pid, []).append(cid)
        # 分配百分比
        for pid in selected_main_chapters:
            sub_ids = group.get(pid, [])
            if not sub_ids:
                continue  # 该大章节下没有被选中的小章节
            sub_ratio = main_ratio / len(sub_ids)
            for cid in sub_ids:
                final_chapter_dist[cid] = sub_ratio
    else:
        # 兼容旧逻辑：全部平均
        avg = 1.0 / len(chapter_ids)
        for cid in chapter_ids:
            final_chapter_dist[int(cid)] = avg

    job_id = uuid.uuid4().hex
    with _jobs_lock:
        _jobs[job_id] = {
            "job_id": job_id,
            "status": "queued",
            "inserted": 0,
            "question_ids": [],
            "created_at": datetime.now().isoformat(timespec="seconds"),
            "seq": 0,
            "events": [],
        }
    app = current_app._get_current_object()
    _executor.submit(_run_generation_v2, app, job_id, int(subject_id), final_chapter_dist, validated_rules, create_user)
    return jsonify({"ok": True, "job_id": job_id, "queued": True})


@ai_bp.post("/generate-answer-sheet-preview")
def generate_answer_sheet_preview():
    payload = request.get_json(silent=True) or {}
    paper_id = payload.get("paper_id")
    if not paper_id:
        return jsonify({"error": {"message": "paper_id 必填", "type": "BadRequest"}}), 400

    try:
        session = get_session(current_app)
        pqr = _table("paper_question_relation")
        qb = _table("question_bank")
        
        stmt = (
            select(qb.c.question_content, qb.c.type_id, qb.c.question_score)
            .join(pqr, pqr.c.question_id == qb.c.question_id)
            .where(pqr.c.paper_id == paper_id)
            .order_by(pqr.c.question_sort)
        )
        questions = session.execute(stmt).mappings().all()
        
        if not questions:
            return jsonify({"error": {"message": "试卷为空", "type": "NotFound"}}), 404
            
        q_text = ""
        for i, q in enumerate(questions, 1):
            q_text += f"{i}. [类型:{q['type_id']}] {q['question_content']} ({q['question_score']}分)\n"
            
        system_prompt = (
            "你是一个专业的答题卡排版助手。请根据提供的试题列表，生成一份符合A3双栏排版标准的Markdown答题卡，不需要题干，只需要答题区域。\n"
            "【整体布局要求】：\n"
            "1. 页面顶部通栏（分栏靠左边）：\n"
            "   - 主标题：居中，字号较大（如<h1>XX考试答题卡</h1>）。\n"
            "   - 考生信息栏：使用表格布局，包含姓名、班级、考号、座位号等填写框。\n"
            "   - 注意事项：列出3-4条填涂规范。\n"
            "   - 准考证号填涂区：右侧生成一个 10列 x 10行 的表格（table class='ticket-no-table'），表头为0-9。\n"
            "2. 答题区域（分两栏）：\n"
            "   - 请将内容包裹在 <div class='columns-container'> ... </div> 中（如果Markdown支持HTML，请直接使用HTML标签以保证布局）。\n"
            "   - 题型之间要有明显的标题（<h2>一、选择题</h2>）。\n"
            "3. 题型排版细节：\n"
            "   - 选择题：每5题一组，使用 <span class='option-box'>A</span> <span class='option-box'>B</span> ... 形式。\n"
            "   - 填空题：生成足够长度的下划线（________________）。\n"
            "   - 简答/计算/作文：生成带有题号的空白区域，可以使用 <div class='essay-line'></div> 重复多次来模拟横线。\n"
            "4. 输出格式：\n"
            "   - 混合使用 Markdown 和 HTML 以达到最佳效果。\n"
            "   - 不要使用代码块包裹，直接输出内容。\n"
            "   - 确保 HTML 标签闭合正确。"
        )
        user_prompt = f"请为以下试题生成答题卡Markdown：\n\n{q_text}"
        
        client = get_deepseek_client()
        
        def stream_response():
            try:
                # First yield the job status (optional, but good for connection check)
                yield f"data: {json.dumps({'type': 'start'})}\n\n"
                
                full_content = ""
                for chunk in client.chat_stream(system_prompt, user_prompt):
                    full_content += chunk
                    # Escape newlines for SSE data payload
                    safe_chunk = json.dumps({"type": "delta", "content": chunk})
                    yield f"data: {safe_chunk}\n\n"
                
                # Final cleanup of the complete content
                markdown_content = full_content
                markdown_content = re.sub(r"^```markdown\s*", "", markdown_content).strip()
                markdown_content = re.sub(r"^```\s*", "", markdown_content).strip()
                markdown_content = re.sub(r"\s*```$", "", markdown_content).strip()
                
                # Send the final cleaned content
                yield f"data: {json.dumps({'type': 'done', 'markdown': markdown_content})}\n\n"
            except Exception as e:
                yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"

        return Response(stream_with_context(stream_response()), mimetype="text/event-stream")
        
    except Exception as e:
        return jsonify({"error": {"message": str(e), "type": "ServerError"}}), 500


def _run_generation_v2(app, job_id: str, subject_id: int, chapter_dist: dict[int, float], rules: list[dict], create_user: str):
    """
    新版生成逻辑：支持章节权重分配
    chapter_dist: {chapter_id: ratio}，ratio Sum 应该约为 1.0
    rules: [{"type_id": 1, "difficulty_id": 1, "count": 10}, ...]
    """
    with app.app_context():
        total_expected = sum(r["count"] for r in rules)
        _job_update(job_id, {"status": "running", "started_at": datetime.now().isoformat(timespec="seconds")})
        _job_event(job_id, "job_start", f"开始生成（共{total_expected}题）", {"total_count": total_expected})
        qb = _table("question_bank")
        ch = _table("textbook_chapter")
        session = get_session(current_app)
        inserted = 0
        created_ids: list[int] = []

        try:
            chapter_ids = list(chapter_dist.keys())
            stmt = select(ch.c.chapter_id, ch.c.chapter_name, ch.c.content, ch.c.parent_chapter_id, ch.c.textbook_id).where(ch.c.chapter_id.in_(chapter_ids))
            chapters_data = {r["chapter_id"]: dict(r) for r in session.execute(stmt).mappings().all()}
            
            if not chapters_data:
                _job_update(job_id, {"status": "error", "error": "所选章节不存在"})
                _job_event(job_id, "job_error", "所选章节不存在")
                return

            # --- 递归获取子章节内容 ---
            # 1. 找出涉及的教材ID
            textbook_ids = set(c["textbook_id"] for c in chapters_data.values() if c.get("textbook_id"))
            
            # 2. 获取这些教材的所有章节以构建层级关系
            all_chapters_map = {} # {chapter_id: row}
            children_map = {} # {parent_id: [child_id, ...]}
            
            if textbook_ids:
                all_ch_stmt = select(ch.c.chapter_id, ch.c.chapter_name, ch.c.parent_chapter_id, ch.c.content).where(ch.c.textbook_id.in_(textbook_ids))
                all_rows = session.execute(all_ch_stmt).mappings().all()
                for row in all_rows:
                    cid = row["chapter_id"]
                    pid = row["parent_chapter_id"]
                    all_chapters_map[cid] = row
                    if pid:
                        children_map.setdefault(pid, []).append(cid)
            
            def get_aggregated_content(root_id):
                contents = []
                # 根节点内容
                root_c = all_chapters_map.get(root_id, {}).get("content")
                if root_c:
                    contents.append(root_c)
                
                # 递归子节点
                stack = [root_id]
                seen = {root_id}
                while stack:
                    curr = stack.pop()
                    children = children_map.get(curr, [])
                    for child in children:
                        if child not in seen:
                            seen.add(child)
                            stack.append(child)
                            c = all_chapters_map.get(child, {}).get("content")
                            if c:
                                contents.append(c)
                return "\n".join(contents)
            # -----------------------

            # 对每条规则进行分配
            for rule_idx, rule in enumerate(rules):
                type_id = rule["type_id"]
                difficulty_id = rule["difficulty_id"]
                total_count = rule["count"]
                
                # 分配数量
                # 使用最大余额法或简单的四舍五入，保证总数一致
                # 这里简单处理：按比例计算，向下取整，剩余的加给占比最大的章节
                
                dist_counts = {} # {chapter_id: count}
                current_sum = 0
                sorted_chapters = sorted(chapter_dist.items(), key=lambda x: x[1], reverse=True)
                
                for cid, ratio in sorted_chapters:
                    c = int(total_count * ratio)
                    dist_counts[cid] = c
                    current_sum += c
                
                # 补齐余数
                remainder = total_count - current_sum
                if remainder > 0:
                    # 加给权重最高的
                    dist_counts[sorted_chapters[0][0]] += remainder
                
                # --- 新增逻辑：将分配给父章节的 count 分摊给所有叶子子章节 ---
                final_tasks = {} # {cid: count}
                
                def get_leaves(node_id):
                    children = children_map.get(node_id, [])
                    if not children:
                        return [node_id]
                    leaves = []
                    for child in children:
                        leaves.extend(get_leaves(child))
                    return leaves

                for cid, count in dist_counts.items():
                    if count <= 0:
                        continue
                    
                    leaves = get_leaves(cid)
                    # 过滤掉不在 all_chapters_map 里的（比如可能没查到）
                    leaves = [l for l in leaves if l in all_chapters_map]
                    
                    if not leaves:
                        # 只有自己
                        final_tasks[cid] = final_tasks.get(cid, 0) + count
                    else:
                        # 平均分配给叶子
                        avg = count // len(leaves)
                        rem = count % len(leaves)
                        for i, leaf in enumerate(leaves):
                            c = avg + (1 if i < rem else 0)
                            if c > 0:
                                final_tasks[leaf] = final_tasks.get(leaf, 0) + c
                
                dist_counts = final_tasks
                # --------------------------------------------------------

                _job_event(
                    job_id,
                    "rule_start",
                    f"规则{rule_idx+1}: 题型={type_id} 难度={difficulty_id} 总数={total_count}",
                    {"type_id": type_id, "difficulty_id": difficulty_id, "total_count": total_count}
                )

                # 逐个章节生成
                for cid, count in dist_counts.items():
                    if count <= 0:
                        continue
                        
                    chapter_info = all_chapters_map.get(cid)
                    if not chapter_info:
                        continue
                        
                    _job_event(job_id, "chapter_start", f"章节：{chapter_info.get('chapter_name')} (分配{count}题)", {"chapter_id": cid})
                    
                    # 复用核心生成逻辑 (call_model 等)
                    # 这里为了避免代码冗余，最好提取一个 _generate_for_single_chapter 函数
                    # 但考虑到 _run_generation 内部有很多闭包变量，暂时内联或者简单提取
                    
                    # --- 开始单章节生成逻辑 ---
                    sample_stmt = (
                        select(qb.c.question_id, qb.c.question_content, qb.c.question_answer, qb.c.question_analysis)
                        .where(and_(qb.c.chapter_id == cid, qb.c.review_status == 1, qb.c.type_id == type_id))
                        .order_by(qb.c.question_id.desc())
                        .limit(3)
                    )
                    sample = session.execute(sample_stmt).mappings().all()
                    source_ids = [str(s["question_id"]) for s in sample]

                    # summary = chapter_info.get("content") or ""
                    # summary = get_aggregated_content(cid)
                    # 修改为只使用当前章节（叶子节点）的内容，避免重复或混淆
                    summary = chapter_info.get("content") or ""
                    
                    system_prompt = "你是出题助理。你必须严格输出 JSON 数组，不要输出任何多余文本。"
                    
                    additional_reqs = ""
                    if type_id == 1: # 单选题
                        additional_reqs = "5) 必须有且只能有4个选项(A/B/C/D)。\n6) 必须有且只能有一个正确答案。\n7) 【强制】题干末尾必须以中文括号（ ）结尾。\n"
                    elif type_id == 7: # 多选题
                        additional_reqs = "5) 必须有且只能有4个选项(A/B/C/D)。\n6) 必须有两个或更多正确答案。\n7) 【强制】题干末尾必须以中文括号（ ）结尾。\n"
                    elif type_id == 2: # 判断题
                        additional_reqs = "5) 题干必须是陈述句。\n6) 答案必须是“正确”或“错误”（或T/F）。\n7) 【强制】题干末尾必须以中文括号（ ）结尾。\n"
                    elif type_id == 4: # 计算题
                        additional_reqs = "5) 必须是计算类题目，禁止出现选项(A/B/C/D)。\n6) 答案需包含具体计算结果。\n"
                    elif type_id == 3: # 填空题
                        additional_reqs = "5) 题干中必须包含填空符（如______）。\n6) 禁止出现选项。\n"
                    elif type_id in [5, 6]: # 简答题, 作文
                        additional_reqs = "5) 必须是主观题，禁止出现选项。\n"

                    content_req = "4) question_content 内可包含选项（如A/B/C/D），但仍是纯文本\n"
                    if type_id not in [1, 7]:
                        content_req = "4) question_content 必须是纯文本，严禁包含选项(A/B/C/D)\n"

                    user_prompt = (
                        "请为指定教材章节生成题目，要求：\n"
                        "1) 只生成与章节相关的题\n"
                        "2) 难度与题型遵循要求\n"
                        "3) 输出严格 JSON 数组，每个元素包含字段：type_id, question_content, question_answer, question_analysis, question_score\n"
                        f"{content_req}"
                        f"{additional_reqs}\n"
                        f"章节名称：{chapter_info['chapter_name']}\n"
                        f"章节概要：{summary if summary else '(暂无概要)'}\n"
                        f"目标数量：{count}\n"
                        f"题型ID：{type_id} (请在返回的JSON中将 type_id 设为 {type_id})\n"
                        f"难度ID：{difficulty_id}\n\n"
                        "参考题目（用于风格与覆盖点，不要重复）：\n"
                    )
                    for i, s in enumerate(sample, start=1):
                        user_prompt += (
                            f"\n{i}. 题干：{s.get('question_content')}\n"
                            f"   答案：{s.get('question_answer')}\n"
                            f"   解析：{s.get('question_analysis')}\n"
                        )

                    client = get_deepseek_client()
                    target_count = int(count)
                    collected: list[dict] = []
                    seen_contents: set[str] = set()
                    attempt = 0

                    def call_model(prompt: str, attempt_no: int, missing: int) -> str:
                        _job_event(
                            job_id,
                            "ai_start",
                            f"请求模型中…（第{attempt_no}次，缺{missing}题）",
                            {"chapter_id": cid, "type_id": type_id, "difficulty_id": difficulty_id, "attempt": attempt_no, "missing": missing},
                        )
                        try:
                            raw_chunks: list[str] = []
                            buf = ""
                            last_flush = time.time()
                            for chunk in client.chat_stream(system_prompt=system_prompt, user_prompt=prompt, temperature=0.7):
                                raw_chunks.append(chunk)
                                buf += chunk
                                now_ts = time.time()
                                if len(buf) >= 200 or "\n" in buf or (now_ts - last_flush) >= 0.8:
                                    _job_event(job_id, "ai_delta", data={"text": buf})
                                    buf = ""
                                    last_flush = now_ts
                            if buf:
                                _job_event(job_id, "ai_delta", data={"text": buf})
                            raw_text = "".join(raw_chunks)
                        except Exception as err:
                            _job_event(job_id, "ai_error", f"流式输出不可用，改用普通请求：{err}")
                            raw_text = client.chat(system_prompt=system_prompt, user_prompt=prompt, temperature=0.7)
                            _job_event(job_id, "ai_delta", data={"text": raw_text})
                        _job_event(job_id, "ai_end", "模型返回完成")
                        return raw_text

                    while len(collected) < target_count and attempt < MAX_FILL_ATTEMPTS:
                        attempt += 1
                        missing = target_count - len(collected)

                        prompt = user_prompt.replace(f"目标数量：{count}\n", f"目标数量：{missing}\n")
                        if attempt > 1:
                            existed = "\n".join([f"- {x['question_content']}" for x in collected[:50]])
                            prompt += (
                                "\n补齐要求：\n"
                                f"1) 你只需要补齐缺口：再生成 {missing} 道新题\n"
                                "2) 只输出 JSON 数组，不要输出 ```json 或任何解释文本\n"
                                "3) 绝对不要重复已有题干\n"
                                f"\n已生成题干（不要重复）：\n{existed}\n"
                            )

                        _job_event(job_id, "rule_retry", f"补齐生成：第{attempt}次（还差{missing}题）")
                        raw = call_model(prompt, attempt, missing)
                        try:
                            items = _extract_json_list(raw)
                        except Exception as err:
                            _job_event(job_id, "rule_warn", f"解析失败（第{attempt}次）：{err}")
                            continue

                        _job_event(job_id, "parse_ok", f"解析成功：{len(items)}条（第{attempt}次）", {"count": len(items), "attempt": attempt})

                        for it in items:
                            if len(collected) >= target_count:
                                break
                            if not isinstance(it, dict):
                                continue
                            content = _pick_first(it, ["question_content", "content", "stem", "question", "题干"])
                            if not content:
                                continue
                            content_key = re.sub(r"\s+", " ", content).strip()
                            if content_key in seen_contents:
                                continue
                            seen_contents.add(content_key)

                            answer = _pick_first(it, ["question_answer", "answer", "答案"])
                            analysis = _pick_first(it, ["question_analysis", "analysis", "解析"])
                            score_raw = _pick_first(it, ["question_score", "score", "分值"])
                            score_val = None
                            if score_raw is not None:
                                try:
                                    score_val = float(score_raw)
                                except Exception:
                                    score_val = None

                            returned_type_id = _pick_first(it, ["type_id", "type", "题型ID"])
                            final_type_id = type_id
                            if returned_type_id:
                                try:
                                    final_type_id = int(returned_type_id)
                                except:
                                    pass

                            collected.append(
                                {
                                    "question_content": content,
                                    "question_answer": answer,
                                    "question_analysis": analysis,
                                    "question_score": score_val,
                                    "type_id": final_type_id,
                                }
                            )

                        _job_event(job_id, "rule_progress", f"已收集：{len(collected)}/{target_count}", {"collected": len(collected), "target": target_count})

                    if len(collected) < target_count:
                        _job_event(job_id, "rule_error", f"补齐失败：需要{target_count}，实际{len(collected)}（已重试{attempt}次）")
                        # 不抛出异常，继续下一个章节
                    
                    now = datetime.now()
                    for it in collected:
                        data = {
                            "subject_id": subject_id,
                            "chapter_id": cid,
                            "type_id": it.get("type_id", type_id),
                            "difficulty_id": difficulty_id,
                            "question_content": it.get("question_content"),
                            "question_answer": it.get("question_answer"),
                            "question_analysis": it.get("question_analysis"),
                            "question_score": it.get("question_score") if it.get("question_score") is not None else 0,
                            "is_ai_generated": 1,
                            "source_question_ids": ",".join(source_ids) if source_ids else None,
                            "review_status": 0,
                            "reviewer": None,
                            "review_time": None,
                            "create_user": create_user,
                            "create_time": now,
                            "update_time": now,
                        }
                        res = session.execute(insert(qb).values(**data))
                        if res.inserted_primary_key:
                            created_ids.append(res.inserted_primary_key[0])
                        inserted += 1
                    
                    _job_update(job_id, {"inserted": inserted, "question_ids": created_ids})
                    _job_event(job_id, "progress", f"已入库：{inserted}题", {"inserted": inserted})
                    # --- 结束单章节生成逻辑 ---

            session.commit()
            _job_update(
                job_id,
                {
                    "status": "done",
                    "inserted": inserted,
                    "question_ids": created_ids,
                    "finished_at": datetime.now().isoformat(timespec="seconds"),
                },
            )
            _job_event(job_id, "job_done", f"任务完成：新增{inserted}题", {"inserted": inserted})
        except Exception as err:
            session.rollback()
            _job_update(
                job_id,
                {
                    "status": "error",
                    "error": str(err),
                    "inserted": inserted,
                    "question_ids": created_ids,
                    "finished_at": datetime.now().isoformat(timespec="seconds"),
                },
            )
            _job_event(job_id, "job_error", str(err))


@ai_bp.get("/jobs/<string:job_id>")
def get_job(job_id: str):
    snap = _job_snapshot(job_id)
    if not snap:
        return jsonify({"error": {"message": "任务不存在", "type": "NotFound"}}), 404
    return jsonify({"job": snap})


@ai_bp.get("/jobs/<string:job_id>/events")
def job_events(job_id: str):
    last_id = request.args.get("last_id", default=0, type=int)

    def gen():
        nonlocal last_id
        yield ": ok\n\n"
        while True:
            snap = _job_snapshot(job_id)
            if not snap:
                yield f"event: error\ndata: {json.dumps({'message': '任务不存在'}, ensure_ascii=False)}\n\n"
                break

            events: list[dict] = []
            with _jobs_lock:
                job = _jobs.get(job_id) or {}
                for e in job.get("events", []):
                    if int(e.get("id") or 0) > last_id:
                        events.append(e)

            for e in events:
                last_id = int(e.get("id") or last_id)
                payload = json.dumps(e, ensure_ascii=False)
                yield f"id: {last_id}\ndata: {payload}\n\n"

            if snap.get("status") in ["done", "error"] and not events:
                break

            time.sleep(0.5)

    headers = {
        "Cache-Control": "no-cache",
        "Connection": "keep-alive",
        "X-Accel-Buffering": "no",
    }
    return Response(stream_with_context(gen()), headers=headers, mimetype="text/event-stream")


@ai_bp.get("/pending")
def list_pending():
    qb = _table("question_bank")
    ch = _table("textbook_chapter")
    sd = _table("subject_dict")
    qtd = _table("question_type_dict")
    qdd = _table("question_difficulty_dict")
    subject_id = request.args.get("subject_id", type=int)
    chapter_ids_str = request.args.get("chapter_id", type=str)
    page = max(1, request.args.get("page", default=1, type=int))
    page_size = min(100, max(1, request.args.get("page_size", default=20, type=int)))

    where = [qb.c.review_status == 0]
    if subject_id is not None:
        where.append(qb.c.subject_id == subject_id)
    if chapter_ids_str:
        try:
            c_ids = [int(x) for x in chapter_ids_str.split(",") if x.strip()]
            if c_ids:
                if len(c_ids) == 1:
                    where.append(qb.c.chapter_id == c_ids[0])
                else:
                    where.append(qb.c.chapter_id.in_(c_ids))
        except Exception:
            pass

    stmt_base = (
        select(
            qb.c.question_id,
            qb.c.subject_id,
            qb.c.chapter_id,
            ch.c.textbook_id,
            qb.c.type_id,
            qb.c.difficulty_id,
            sd.c.subject_name.label("subject_name"),
            ch.c.chapter_name.label("chapter_name"),
            qtd.c.type_name.label("type_name"),
            qdd.c.difficulty_name.label("difficulty_name"),
            qb.c.question_content,
            qb.c.question_answer,
            qb.c.question_analysis,
            qb.c.question_score,
            qb.c.is_ai_generated,
            qb.c.source_question_ids,
            qb.c.create_user,
            qb.c.create_time,
        )
        .select_from(
            qb.outerjoin(ch, ch.c.chapter_id == qb.c.chapter_id)
            .outerjoin(sd, sd.c.subject_id == qb.c.subject_id)
            .outerjoin(qtd, qtd.c.type_id == qb.c.type_id)
            .outerjoin(qdd, qdd.c.difficulty_id == qb.c.difficulty_id)
        )
        .where(and_(*where))
    )
    
    stmt = (
        stmt_base
        .order_by(qb.c.question_id.desc())
        .offset((page - 1) * page_size)
        .limit(page_size)
    )

    try:
        # 使用不带分页的查询来计算总数
        count_query = select(func.count()).select_from(stmt_base.subquery())
        total = get_session(current_app).execute(count_query).scalar_one()
        
        rows = get_session(current_app).execute(stmt).mappings().all()
        return jsonify({"items": [dict(r) for r in rows], "page": page, "page_size": page_size, "total": total})
    except SQLAlchemyError as err:
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


@ai_bp.post("/generate-smart-paper")
def generate_smart_paper():
    payload = request.get_json(silent=True) or {}
    subject_id = payload.get("subject_id")
    textbook_id = payload.get("textbook_id")
    description = payload.get("description")
    create_user = payload.get("create_user") or "ai"

    if not subject_id or not textbook_id or not description:
        return jsonify({"error": {"message": "subject_id, textbook_id, description 必填", "type": "BadRequest"}}), 400

    job_id = uuid.uuid4().hex
    with _jobs_lock:
        _jobs[job_id] = {
            "job_id": job_id,
            "status": "queued",
            "type": "smart_paper",
            "create_user": create_user,
            "created_at": datetime.now().isoformat(timespec="seconds"),
            "seq": 0,
            "events": [],
        }
    
    app = current_app._get_current_object()
    _executor.submit(_do_smart_paper_job, app, job_id, subject_id, textbook_id, description)

    return jsonify({"job_id": job_id})

def _do_smart_paper_job(app, job_id, subject_id, textbook_id, description):
    with app.app_context():
        try:
            _job_update(job_id, {"status": "running", "started_at": datetime.now().isoformat(timespec="seconds")})
            _job_event(job_id, "job_start", "开始智能组卷分析...")
            session = get_session(app)
            
            # 1. Fetch Metadata
            _job_event(job_id, "meta_fetch", "正在获取基础数据...")
            
            # Types
            qtd = _table("question_type_dict")
            types = session.execute(select(qtd.c.type_id, qtd.c.type_name)).all()
            type_info = ", ".join([f"{t.type_name}(ID={t.type_id})" for t in types])

            # Difficulties
            qdd = _table("question_difficulty_dict")
            diffs = session.execute(select(qdd.c.difficulty_id, qdd.c.difficulty_name)).all()
            diff_info = ", ".join([f"{d.difficulty_name}(ID={d.difficulty_id})" for d in diffs])

            # Chapters
            ch = _table("textbook_chapter")
            chapters = session.execute(select(ch.c.chapter_id, ch.c.chapter_name).where(ch.c.textbook_id == textbook_id)).all()
            chapter_info = "\n".join([f"- {c.chapter_name} (ID={c.chapter_id})" for c in chapters])

            # 2. AI Analysis
            _job_event(job_id, "ai_analyze", "正在分析您的需求...")
            client = get_deepseek_client()
            
            system_prompt = (
                "你是一个专业的试卷生成助手。你的任务是根据用户的自然语言描述，分析出试卷的结构和题目要求。\\n"
                "你需要返回一个严格的JSON对象，包含试卷的基本信息和题目列表。\\n"
                "请严格按照以下JSON格式输出，不要输出任何Markdown标记或解释文字：\\n"
                "{\\n"
                "  \"paper_name\": \"试卷名称\",\\n"
                "  \"exam_duration\": 120, // 考试时长(分钟)\\n"
                "  \"is_closed_book\": true, // 是否闭卷\\n"
                "  \"paper_desc\": \"试卷描述\",\\n"
                "  \"sections\": [ // 试卷大题/部分\\n"
                "    {\\n"
                "      \"name\": \"一、选择题\", // 大题名称\\n"
                "      \"type_id\": 1, // 题型ID\\n"
                "      \"difficulty_id\": 1, // 难度ID (可选，如果用户指定了整体难度或该题型难度)\\n"
                "      \"count\": 10, // 题目数量\\n"
                "      \"score_per_question\": 2.0, // 每题分值\\n"
                "      \"chapter_ids\": [1, 2], // 涉及章节ID列表 (可选，为空则从整本教材随机)\\n"
                "      \"keywords\": [\"关键词1\"] // 关键词 (可选，用于模糊匹配题干)\\n"
                "    }\\n"
                "  ]\\n"
                "}"
            )

            user_prompt = (
                f"【环境信息】\\n"
                f"可用题型：{type_info}\\n"
                f"可用难度：{diff_info}\\n"
                f"可用章节：\\n{chapter_info}\\n\\n"
                f"【用户需求】\\n{description}\\n\\n"
                f"请根据用户需求生成试卷结构JSON。"
            )

            _job_event(job_id, "ai_thinking", "AI正在思考组卷策略...")
            
            # Call AI
            raw_response = ""
            # Stream the AI response text to frontend for "thinking" effect
            buf = ""
            last_flush = time.time()
            for chunk in client.chat_stream(system_prompt=system_prompt, user_prompt=user_prompt, temperature=0.7):
                raw_response += chunk
                buf += chunk
                now_ts = time.time()
                if len(buf) >= 100 or "\\n" in buf or (now_ts - last_flush) >= 0.5:
                    _job_event(job_id, "ai_delta", data={"text": buf})
                    buf = ""
                    last_flush = now_ts
            
            if buf:
                _job_event(job_id, "ai_delta", data={"text": buf})

            _job_event(job_id, "ai_parsed", "策略生成完成，正在解析...")
            
            plan = None
            try:
                # Try strict JSON parsing first
                start = raw_response.find("{")
                end = raw_response.rfind("}")
                if start >= 0 and end >= 0:
                     candidate = raw_response[start : end + 1]
                     plan = json.loads(candidate)
            except Exception:
                pass
            
            if not plan:
                 _job_event(job_id, "job_error", "AI未能生成有效的JSON策略")
                 return

            # 3. Execute Query & Assemble Paper
            _job_event(job_id, "db_query", "正在根据策略抽取题目...")
            
            qb = _table("question_bank")
            picked_questions = []
            
            total_score = 0
            
            # Pre-fetch all chapters if needed for filtering
            all_chapter_ids = [c.chapter_id for c in chapters]

            for section in plan.get("sections", []):
                type_id = section.get("type_id")
                diff_id = section.get("difficulty_id")
                count = section.get("count", 0)
                score = section.get("score_per_question", 0)
                c_ids = section.get("chapter_ids", [])
                keywords = section.get("keywords", [])
                
                if count <= 0:
                    continue
                
                # Build Query
                conditions = [
                    qb.c.subject_id == subject_id,
                    qb.c.type_id == type_id,
                    qb.c.review_status == 1 # Only reviewed questions
                ]
                
                if diff_id:
                    conditions.append(qb.c.difficulty_id == diff_id)
                
                if c_ids:
                    # Validate chapter IDs
                    valid_cids = [cid for cid in c_ids if cid in all_chapter_ids]
                    if valid_cids:
                         conditions.append(qb.c.chapter_id.in_(valid_cids))
                    else:
                         conditions.append(qb.c.chapter_id.in_(all_chapter_ids))
                else:
                    conditions.append(qb.c.chapter_id.in_(all_chapter_ids))

                if keywords:
                    keyword_conditions = [qb.c.question_content.like(f"%{k}%") for k in keywords]
                    if keyword_conditions:
                        conditions.append(or_(*keyword_conditions))

                stmt = select(
                    qb.c.question_id, 
                    qb.c.question_content, 
                    qb.c.question_answer, 
                    qb.c.question_analysis,
                    qb.c.type_id,
                    qb.c.difficulty_id,
                    qb.c.chapter_id,
                    qb.c.subject_id
                ).where(and_(*conditions))
                
                # Randomize
                stmt = stmt.order_by(func.random()).limit(count)
                
                rows = session.execute(stmt).mappings().all()
                
                for row in rows:
                    q = dict(row)
                    q["question_score"] = score # Override score based on plan
                    # Add section info if useful
                    picked_questions.append(q)
                    total_score += score
                
                if len(rows) < count:
                    _job_event(job_id, "warn", f"题型ID={type_id} 数量不足，需求{count}，实际找到{len(rows)}")

            # 4. Finish
            result = {
                "paper_name": plan.get("paper_name", "AI生成试卷"),
                "paper_desc": plan.get("paper_desc", ""),
                "exam_duration": plan.get("exam_duration", 120),
                "is_closed_book": plan.get("is_closed_book", False),
                "questions": picked_questions
            }
            
            _job_update(job_id, {"status": "done", "finished_at": datetime.now().isoformat(timespec="seconds"), "result": result})
            _job_event(job_id, "job_done", f"组卷完成，共{len(picked_questions)}题", result)

        except Exception as e:
            import traceback
            traceback.print_exc()
            _job_update(job_id, {"status": "error", "error": str(e)})
            _job_event(job_id, "job_error", str(e))
        finally:
             close_session()


@ai_bp.post("/verify")
def verify_question():
    payload = request.get_json(silent=True) or {}
    question_id = payload.get("question_id")
    action = payload.get("action")
    reviewer = payload.get("reviewer") or "reviewer"
    fields = payload.get("fields") or {}

    if not question_id or action not in ["approve", "reject", "update_and_approve"]:
        return jsonify({"error": {"message": "question_id 与 action 必填", "type": "BadRequest"}}), 400

    qb = _table("question_bank")
    now = datetime.now()

    data = {"reviewer": reviewer, "review_time": now}
    if action == "reject":
        data["review_status"] = 2
    else:
        data["review_status"] = 1
        if action == "update_and_approve":
            for k in ["question_content", "question_answer", "question_analysis", "question_score", "type_id", "difficulty_id", "chapter_id", "subject_id"]:
                if k in fields:
                    data[k] = fields.get(k)

    try:
        session = get_session(current_app)
        session.execute(update(qb).where(qb.c.question_id == int(question_id)).values(**data))
        session.commit()
        return jsonify({"ok": True})
    except SQLAlchemyError as err:
        get_session(current_app).rollback()
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500


def _run_parsing(app, job_id: str, text: str, subject_id: Optional[int], chapter_id: Optional[int], type_id: Optional[int], difficulty_id: Optional[int], create_user: str):
    with app.app_context():
        _job_update(job_id, {"status": "running", "started_at": datetime.now().isoformat(timespec="seconds")})
        
        # Estimate total count
        estimated_count = 0
        try:
            matches = re.findall(r"^\d+[\.、\s]", text, re.MULTILINE)
            if matches:
                estimated_count = len(matches)
        except:
            pass
            
        _job_event(job_id, "job_start", f"开始AI解析（预计{estimated_count}题）", {"total_count": estimated_count})
        
        # qb = _table("question_bank")
        # session = get_session(current_app)
        # inserted = 0
        # created_ids: list[int] = []
        parsed_items: list[dict] = []
        
        try:
            system_prompt = "你是出题助理。你必须严格输出 JSON 数组，不要输出任何多余文本。"
            user_prompt = (
                "请解析以下题目文本，提取题目信息。\n"
                "要求：\n"
                "1) 输出严格 JSON 数组，每个元素包含：type_id, question_content, question_answer, question_analysis, question_score\n"
                "2) 如果文本中缺少答案或解析，请你根据题目内容自动生成正确的答案和详细解析\n"
                "3) question_content 应包含题干和选项（如果是选择题）\n"
                "4) 自动判断题目分值，如果无法判断则默认为 2\n"
                "5) 自动判断题型并返回 type_id：1=单选题, 2=判断题, 3=填空题, 4=计算题, 5=简答题, 6=作文题, 7=多选题。如果无法判断则默认为 5\n"
                "\n待解析文本：\n"
                f"{text[:50000]}"
            )
            
            client = get_deepseek_client()
            
            _job_event(job_id, "ai_start", "正在请求AI进行解析...")
            
            raw_chunks: list[str] = []
            buf = ""
            last_flush = time.time()
            
            try:
                for chunk in client.chat_stream(system_prompt=system_prompt, user_prompt=user_prompt, temperature=0.2):
                    raw_chunks.append(chunk)
                    buf += chunk
                    now_ts = time.time()
                    if len(buf) >= 200 or "\n" in buf or (now_ts - last_flush) >= 0.8:
                        _job_event(job_id, "ai_delta", data={"text": buf})
                        buf = ""
                        last_flush = now_ts
                if buf:
                    _job_event(job_id, "ai_delta", data={"text": buf})
                raw_text = "".join(raw_chunks)
            except Exception as err:
                 _job_event(job_id, "ai_error", f"流式请求失败: {err}")
                 raw_text = client.chat(system_prompt=system_prompt, user_prompt=user_prompt, temperature=0.2)
                 _job_event(job_id, "ai_delta", data={"text": raw_text})
            
            _job_event(job_id, "ai_end", "AI响应完成，开始提取数据")
            
            items = []
            try:
                items = _extract_json_list(raw_text)
            except Exception as err:
                _job_event(job_id, "job_error", f"解析JSON失败: {err}")
                raise
            
            _job_event(job_id, "parse_ok", f"成功解析出 {len(items)} 道题目")
            
            now = datetime.now()
            for it in items:
                if not isinstance(it, dict): continue
                
                content = _pick_first(it, ["question_content", "content", "stem", "question", "题干"])
                if not content: continue
                
                answer = _pick_first(it, ["question_answer", "answer", "答案"])
                analysis = _pick_first(it, ["question_analysis", "analysis", "解析"])
                score_raw = _pick_first(it, ["question_score", "score", "分值"])
                score_val = None
                if score_raw:
                    try:
                        score_val = float(score_raw)
                    except:
                        score_val = 2.0
                else:
                    score_val = 2.0

                # Determine type_id: detected > provided > default(5)
                final_type_id = type_id
                detected_type = it.get("type_id")
                if detected_type:
                    try:
                        dt = int(detected_type)
                        if dt > 0:
                            final_type_id = dt
                    except:
                        pass
                
                if not final_type_id:
                    final_type_id = 5

                data = {
                    "subject_id": subject_id,
                    "chapter_id": chapter_id,
                    "type_id": final_type_id,
                    "difficulty_id": difficulty_id,
                    "question_content": content,
                    "question_answer": answer,
                    "question_analysis": analysis,
                    "question_score": score_val,
                    "is_ai_generated": 1,
                    "source_question_ids": None,
                    "review_status": 0,
                    # "reviewer": None,
                    # "review_time": None,
                    "create_user": create_user,
                    # "create_time": now,
                    # "update_time": now,
                }
                parsed_items.append(data)
                # res = session.execute(insert(qb).values(**data))
                # if res.inserted_primary_key:
                #     created_ids.append(res.inserted_primary_key[0])
                # inserted += 1
            
            # session.commit()
            _job_update(
                job_id,
                {
                    "status": "done",
                    "inserted": 0,
                    "items": parsed_items,
                    "question_ids": [],
                    "finished_at": datetime.now().isoformat(timespec="seconds"),
                },
            )
            _job_event(job_id, "job_done", f"解析完成：共{len(parsed_items)}题", {"count": len(parsed_items)})
            
        except Exception as err:
            # session.rollback()
            _job_update(job_id, {"status": "error", "error": str(err)})
            _job_event(job_id, "job_error", str(err))


@ai_bp.post("/parse-word")
def parse_word():
    file = request.files.get("file")
    if file is None:
        return jsonify({"error": {"message": "缺少上传文件 file", "type": "BadRequest"}}), 400
        
    filename = file.filename or ""
    lower_name = filename.lower()
    
    subject_id = request.form.get("subject_id", type=int)
    chapter_id = request.form.get("chapter_id", type=int)
    type_id = request.form.get("type_id", type=int)
    difficulty_id = request.form.get("difficulty_id", type=int)
    create_user = request.form.get("create_user") or "ai_import"

    # 移除必填校验，因为可以在解析后手动设置
    # if not subject_id or not chapter_id or not type_id or not difficulty_id:
    #     return jsonify({"error": {"message": "subject_id, chapter_id, type_id, difficulty_id 必填", "type": "BadRequest"}}), 400
    
    text = ""
    try:
        file_bytes = file.read()
        if lower_name.endswith(".pdf"):
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                pages_text = []
                for page in pdf.pages:
                    pt = page.extract_text()
                    if pt:
                        pages_text.append(pt)
                text = "\n".join(pages_text)
        else:
            # 默认尝试 Word
            doc = Document(io.BytesIO(file_bytes))
            text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
    except Exception as e:
        return jsonify({"error": {"message": f"读取文件失败: {str(e)}", "type": "BadRequest"}}), 400
        
    if not text:
        return jsonify({"error": {"message": "文件内容为空或无法提取文本", "type": "BadRequest"}}), 400
        
    job_id = uuid.uuid4().hex
    with _jobs_lock:
        _jobs[job_id] = {
            "job_id": job_id,
            "status": "queued",
            "inserted": 0,
            "items": [], # 存储解析结果
            "question_ids": [],
            "created_at": datetime.now().isoformat(timespec="seconds"),
            "seq": 0,
            "events": [],
        }
        
    app = current_app._get_current_object()
    _executor.submit(_run_parsing, app, job_id, text, subject_id, chapter_id, type_id, difficulty_id, create_user)
    
    return jsonify({"ok": True, "job_id": job_id, "queued": True})


@ai_bp.post("/verify/batch")
def verify_batch():
    payload = request.get_json(silent=True) or {}
    action = payload.get("action")
    reviewer = payload.get("reviewer") or "reviewer"
    ids = payload.get("ids") or []
    
    qb = _table("question_bank")
    now = datetime.now()
    
    if action == "approve_all_pending":
        stmt = (
            update(qb)
            .where(qb.c.review_status == 0)
            .values(review_status=1, reviewer=reviewer, review_time=now)
        )
    elif action == "approve_selected":
        if not ids:
             return jsonify({"error": {"message": "ids 不能为空", "type": "BadRequest"}}), 400
        stmt = (
            update(qb)
            .where(qb.c.question_id.in_(ids))
            .values(review_status=1, reviewer=reviewer, review_time=now)
        )
    else:
        return jsonify({"error": {"message": "不支持的操作", "type": "BadRequest"}}), 400
    
    try:
        session = get_session(current_app)
        res = session.execute(stmt)
        session.commit()
        return jsonify({"ok": True, "count": res.rowcount})
    except SQLAlchemyError as err:
        get_session(current_app).rollback()
        return jsonify({"error": {"message": str(err), "type": err.__class__.__name__}}), 500
