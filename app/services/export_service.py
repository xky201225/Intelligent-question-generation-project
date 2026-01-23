import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from fpdf import FPDF
from flask import current_app

class CustomPDF(FPDF):
    def __init__(self, paper, questions, is_answer=False, font_path=None):
        super().__init__('P', 'mm', 'A4')
        self.paper = paper
        self.questions = questions
        self.is_answer = is_answer
        self.font_path = font_path
        self.set_margins(20, 25, 20)
        
        # Load font
        if self.font_path and os.path.exists(self.font_path):
            try:
                # Add font (requires fpdf2 or recent fpdf with unicode support)
                # Note: 'uni=True' is for old fpdf, fpdf2 handles it differently.
                # Assuming standard fpdf 1.7.2 as requested, we use add_font with uni=True if supported
                # or just standard add_font for TTF.
                self.add_font('KaiTi', '', self.font_path, uni=True)
                self.set_font('KaiTi', '', 12)
            except Exception as e:
                print(f"Error loading font: {e}")
                self.set_font('Arial', '', 12)
        else:
            self.set_font('Arial', '', 12)

    def header(self):
        pass

    def footer(self):
        self.set_y(-15)
        if 'KaiTi' in self.fonts:
            self.set_font('KaiTi', '', 8)
        else:
            self.set_font('Arial', '', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

    def generate_content(self):
        # Title
        if 'KaiTi' in self.fonts:
            self.set_font('KaiTi', '', 16)
        else:
            self.set_font('Arial', 'B', 16)
            
        title = f"{self.paper['paper_name']} - Answer Key" if self.is_answer else self.paper['paper_name']
        self.cell(0, 15, title, 0, 1, 'C')
        
        # Info
        if 'KaiTi' in self.fonts:
            self.set_font('KaiTi', '', 10)
        else:
            self.set_font('Arial', '', 10)
            
        self.cell(0, 8, f"Subject: {self.paper['subject_name']}  Score: {self.paper['total_score']}  Creator: {self.paper['creator']}", 0, 1, 'L')
        self.ln(10)
        
        # Content
        for idx, q in enumerate(self.questions, 1):
            if 'KaiTi' in self.fonts:
                self.set_font('KaiTi', '', 11)
            else:
                self.set_font('Arial', '', 11)
                
            # Question Stem
            self.multi_cell(0, 8, f"{idx}. {q['content']} ({q['score']} pts)", 0, 'L')
            self.ln(2)
            
            # Options (if any, assuming encoded in content or we need to parse it)
            # The current QuestionBank model stores content as text. If options are separate, they need handling.
            # In the 1/ code, options are stored in JSON. In my current model, they are likely part of content text 
            # or I need to update my model. For now, assuming content has it.
            
            if self.is_answer:
                self.ln(2)
                self.multi_cell(0, 8, f"Answer: {q['answer']}", 0, 'L')
                self.multi_cell(0, 8, f"Analysis: {q['analysis']}", 0, 'L')
                
            self.ln(5)

class ExportService:
    @staticmethod
    def _get_font_path():
        # Try to find a Chinese font
        # Windows fonts
        paths = [
            r"C:\Windows\Fonts\simkai.ttf", # KaiTi
            r"C:\Windows\Fonts\simhei.ttf", # SimHei
            r"C:\Windows\Fonts\msyh.ttf",   # Microsoft YaHei
            os.path.join(current_app.root_path, 'static', 'fonts', 'simkai.ttf')
        ]
        for p in paths:
            if os.path.exists(p):
                return p
        return None

    @staticmethod
    def export_to_word(paper_data, questions, with_answers=False):
        doc = Document()
        
        # Set style to support Chinese
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial' # Fallback
        font.size = Pt(12)
        r = style._element.rPr.rFonts
        r.set(qn('w:eastAsia'), 'KaiTi') # Try setting East Asia font
        
        doc.add_heading(paper_data['paper_name'], 0)
        doc.add_paragraph(f"科目: {paper_data['subject_name']}  总分: {paper_data['total_score']}  出卷人: {paper_data['creator']}")
        doc.add_paragraph("-" * 50)
        
        for idx, q in enumerate(questions, 1):
            p = doc.add_paragraph()
            p.add_run(f"第{idx}题 ({q['score']}分): ").bold = True
            p.add_run(q['content'])
            
            if with_answers:
                p2 = doc.add_paragraph()
                p2.add_run(f"答案: {q['answer']}").italic = True
                doc.add_paragraph(f"解析: {q['analysis']}")
            
            doc.add_paragraph("") # Space
            
        filename = f"{paper_data['paper_name']}_{'教师版' if with_answers else '学生版'}.docx"
        filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], 'exports', filename)
        
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        doc.save(filepath)
        return filepath

    @staticmethod
    def export_to_pdf(paper_data, questions, with_answers=False):
        filename = f"{paper_data['paper_name']}_{'教师版' if with_answers else '学生版'}.pdf"
        filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], 'exports', filename)
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        
        font_path = ExportService._get_font_path()
        
        pdf = CustomPDF(paper_data, questions, is_answer=with_answers, font_path=font_path)
        pdf.add_page()
        pdf.generate_content()
        
        pdf.output(filepath, 'F')
        return filepath
